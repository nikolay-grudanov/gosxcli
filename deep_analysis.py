#!/usr/bin/env python3
"""Extract complete XML from DOCX template."""

import sys
sys.path.insert(0, '/home/gna/.local/lib/python3.14/site-packages')

from docx import Document
from docx.oxml.ns import qn
from lxml import etree
import zipfile

TEMPLATE = "/home/gna/workspase/projects/gosxcli/template/Шаблон_оформления_ВКР_2026_новый.docx"

# First extract raw XML files from ZIP
print("FILES IN DOCX:")
print("=" * 80)
with zipfile.ZipFile(TEMPLATE) as z:
    for name in sorted(z.namelist()):
        info = z.getinfo(name)
        print(f"  {name} ({info.file_size} bytes)")
    
    print("\n\n")
    print("=" * 80)
    print("STYLES.XML (COMPLETE)")
    print("=" * 80)
    styles_xml = z.read('word/styles.xml')
    root = etree.fromstring(styles_xml)
    print(etree.tostring(root, pretty_print=True).decode())
    
    print("\n\n")
    print("=" * 80)
    print("NUMBERING.XML (COMPLETE)")
    print("=" * 80)
    if 'word/numbering.xml' in z.namelist():
        num_xml = z.read('word/numbering.xml')
        num_root = etree.fromstring(num_xml)
        print(etree.tostring(num_root, pretty_print=True).decode())
    else:
        print("NOT FOUND")
    
    print("\n\n")
    print("=" * 80)
    print("DOCUMENT.XML.RELS (COMPLETE)")
    print("=" * 80)
    rels_xml = z.read('word/_rels/document.xml.rels')
    rels_root = etree.fromstring(rels_xml)
    print(etree.tostring(rels_root, pretty_print=True).decode())
    
    print("\n\n")
    print("=" * 80)
    print("SETTINGS.XML (COMPLETE)")
    print("=" * 80)
    settings_xml = z.read('word/settings.xml')
    settings_root = etree.fromstring(settings_xml)
    print(etree.tostring(settings_root, pretty_print=True).decode())

    print("\n\n")
    print("=" * 80)
    print("CONTENT_TYPES.XML (COMPLETE)")
    print("=" * 80)
    ct_xml = z.read('[Content_Types].xml')
    ct_root = etree.fromstring(ct_xml)
    print(etree.tostring(ct_root, pretty_print=True).decode())

# Now analyze specific document elements using python-docx
doc = Document(TEMPLATE)

print("\n\n")
print("=" * 80)
print("ALL PARAGRAPHS WITH DETAILS")
print("=" * 80)

for i, para in enumerate(doc.paragraphs):
    text = para.text[:120] if para.text else ""
    style_name = para.style.name if para.style else "None"
    
    # Only print paragraphs with content or non-Normal style
    if text.strip() or style_name != 'Normal':
        pf = para.paragraph_format
        
        alignment = para.alignment
        runs_info = []
        for run in para.runs:
            ri = f"text='{run.text[:60]}'"
            if run.bold: ri += " BOLD"
            if run.italic: ri += " ITALIC"
            if run.underline: ri += f" UL({run.underline})"
            if run.font.size: ri += f" {run.font.size.pt}pt"
            if run.font.name: ri += f" font='{run.font.name}'"
            if run.font.color and run.font.color.rgb:
                ri += f" color=#{run.font.color.rgb}"
            runs_info.append(ri)
        
        print(f"\nP{i:03d}: style='{style_name}' align={alignment}")
        if text:
            print(f"  text: '{text}'")
        for ri in runs_info:
            print(f"  run: {ri}")
        
        # Print paragraph XML properties for important styles
        if 'Heading' in style_name or 'Титул' in style_name or 'Подзаголовок' in style_name or 'Подпись' in style_name or 'Таб��ица' in style_name or 'Формул' in style_name or 'toc' in style_name or 'TOC' in style_name or 'Заг_' in style_name or 'List' in style_name:
            pPr = para._element.find(qn('w:pPr'))
            if pPr is not None:
                print(f"  pPr XML: {etree.tostring(pPr, pretty_print=False).decode()}")

print("\n\n")
print("=" * 80)
print("ALL TABLES WITH COMPLETE DETAILS")
print("=" * 80)

for ti, table in enumerate(doc.tables):
    print(f"\n{'━' * 60}")
    print(f"TABLE {ti}: {len(table.rows)} rows x {len(table.columns)} cols")
    print(f"  style: {table.style.name if table.style else 'None'}")
    
    # Full table properties XML
    tbl = table._tbl
    tblPr = tbl.find(qn('w:tblPr'))
    if tblPr is not None:
        print(f"  tblPr XML:")
        print(etree.tostring(tblPr, pretty_print=True).decode())
    
    # Table grid (column widths)
    tblGrid = tbl.find(qn('w:tblGrid'))
    if tblGrid is not None:
        print(f"  tblGrid XML: {etree.tostring(tblGrid, pretty_print=False).decode()}")
    
    # Each row and cell
    for ri, row in enumerate(table.rows):
        print(f"\n  Row {ri}:")
        trPr = row._tr.find(qn('w:trPr'))
        if trPr is not None:
            print(f"    trPr: {etree.tostring(trPr, pretty_print=False).decode()}")
        
        for ci, cell in enumerate(row.cells):
            text = cell.text[:100]
            print(f"    Cell[{ri},{ci}]: '{text}'")
            
            # Cell properties
            tc = cell._tc
            tcPr = tc.find(qn('w:tcPr'))
            if tcPr is not None:
                print(f"      tcPr: {etree.tostring(tcPr, pretty_print=False).decode()}")
            
            # Cell paragraph formatting
            for pi, p in enumerate(cell.paragraphs):
                if p.text.strip():
                    print(f"      para[{pi}]: style='{p.style.name}' text='{p.text[:80]}'")
                    for r in p.runs:
                        ri_info = f"'{r.text[:50]}'"
                        if r.bold: ri_info += " B"
                        if r.italic: ri_info += " I"
                        if r.font.size: ri_info += f" {r.font.size.pt}pt"
                        if r.font.name: ri_info += f" f={r.font.name}"
                        print(f"        run: {ri_info}")

print("\n\n")
print("=" * 80)
print("HEADERS AND FOOTERS")
print("=" * 80)

for si, section in enumerate(doc.sections):
    print(f"\n--- Section {si} ---")
    
    for hdr_type, hdr in [("Header", section.header), ("Footer", section.footer)]:
        if hdr is not None:
            print(f"\n  {hdr_type}:")
            print(f"    is_linked_to_previous: {hdr.is_linked_to_previous}")
            for pi, p in enumerate(hdr.paragraphs):
                print(f"    para[{pi}]: style='{p.style.name}' text='{p.text}'")
                for r in p.runs:
                    print(f"      run: '{r.text}' bold={r.bold} font={r.font.name} size={r.font.size}")
                # Check for fields
                fld_chars = p._element.findall('.//' + qn('w:fldChar'))
                instr_texts = p._element.findall('.//' + qn('w:instrText'))
                for fc in fld_chars:
                    print(f"      fldChar: type={fc.get(qn('w:fldCharType'))}")
                for it in instr_texts:
                    print(f"      instrText: '{it.text}'")
            
            # Header/footer XML
            print(f"    Full XML:")
            print(etree.tostring(hdr._element, pretty_print=True).decode()[:2000])

print("\n\n")
print("=" * 80)
print("IMAGE DETAILS")
print("=" * 80)

for i, para in enumerate(doc.paragraphs):
    for run in para.runs:
        drawings = run._element.findall(qn('w:drawing'))
        for d_idx, d in enumerate(drawings):
            print(f"\nImage in P{i}, run='{run.text[:30]}', style={para.style.name}")
            xml_str = etree.tostring(d, pretty_print=True).decode()
            print(f"Drawing XML:\n{xml_str[:3000]}")

print("\nDONE")