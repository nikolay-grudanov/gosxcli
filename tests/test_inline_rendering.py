"""Test inline formatting rendering."""

from __future__ import annotations

import tempfile
from pathlib import Path

from docx import Document as DocxDocument

from typst_gost_docx.ir.model import (
    Document,
    InlineCodeNode,
    InlineRunNode,
    NodeType,
    Paragraph,
    TextRun,
)
from typst_gost_docx.writers.docx_writer import DocxWriter


def test_inline_bold_rendering() -> None:
    """Test bold text rendering."""
    # Create IR document with bold text
    bold_run = InlineRunNode(
        node_type=NodeType.INLINE_RUN,
        id="test-1",
        text="bold",
        bold=True,
    )
    paragraph = Paragraph(
        id="test-para",
        runs=[
            TextRun(node_type=NodeType.TEXT_RUN, id="test-2", text="This is "),
            bold_run,
        ],
    )
    doc = Document(id="test-doc", blocks=[paragraph])

    # Write to DOCX
    with tempfile.TemporaryDirectory() as tmpdir:
        output_path = Path(tmpdir) / "output.docx"
        writer = DocxWriter()
        writer.write(doc, output_path)

        # Verify DOCX content
        docx_doc = DocxDocument(str(output_path))
        para = docx_doc.paragraphs[0]
        runs = para.runs

        # Find the bold run
        bold_runs = [r for r in runs if r.bold]
        assert len(bold_runs) == 1, f"Expected 1 bold run, found {len(bold_runs)}"
        assert bold_runs[0].text == "bold", f"Expected 'bold', got '{bold_runs[0].text}'"


def test_inline_italic_rendering() -> None:
    """Test italic text rendering."""
    # Create IR document with italic text
    italic_run = InlineRunNode(
        node_type=NodeType.INLINE_RUN,
        id="test-1",
        text="italic",
        italic=True,
    )
    paragraph = Paragraph(
        id="test-para",
        runs=[
            TextRun(node_type=NodeType.TEXT_RUN, id="test-2", text="This is "),
            italic_run,
        ],
    )
    doc = Document(id="test-doc", blocks=[paragraph])

    # Write to DOCX
    with tempfile.TemporaryDirectory() as tmpdir:
        output_path = Path(tmpdir) / "output.docx"
        writer = DocxWriter()
        writer.write(doc, output_path)

        # Verify DOCX content
        docx_doc = DocxDocument(str(output_path))
        para = docx_doc.paragraphs[0]
        runs = para.runs

        # Find the italic run
        italic_runs = [r for r in runs if r.italic]
        assert len(italic_runs) == 1, f"Expected 1 italic run, found {len(italic_runs)}"
        assert italic_runs[0].text == "italic", f"Expected 'italic', got '{italic_runs[0].text}'"


def test_inline_code_rendering() -> None:
    """Test code text rendering with monospace font."""
    # Create IR document with code text
    code_node = InlineCodeNode(
        node_type=NodeType.INLINE_CODE,
        id="test-1",
        code="print('hello')",
    )
    paragraph = Paragraph(
        id="test-para",
        runs=[
            TextRun(node_type=NodeType.TEXT_RUN, id="test-2", text="Code: "),
            code_node,
        ],
    )
    doc = Document(id="test-doc", blocks=[paragraph])

    # Write to DOCX
    with tempfile.TemporaryDirectory() as tmpdir:
        output_path = Path(tmpdir) / "output.docx"
        writer = DocxWriter()
        writer.write(doc, output_path)

        # Verify DOCX content
        docx_doc = DocxDocument(str(output_path))
        para = docx_doc.paragraphs[0]
        runs = para.runs

        # Find the code run (should have Courier font or Code style)
        code_runs = [r for r in runs if r.text == "print('hello')"]
        assert len(code_runs) == 1, f"Expected 1 code run, found {len(code_runs)}"

        # Check font is monospace (Courier New)
        code_run = code_runs[0]
        font_name = code_run.font.name
        assert font_name and "Courier" in font_name, f"Expected Courier font, got '{font_name}'"


def test_inline_mixed_rendering() -> None:
    """Test mixed formatting in one paragraph."""
    # Create IR document with mixed formatting
    bold_run = InlineRunNode(
        node_type=NodeType.INLINE_RUN,
        id="test-1",
        text="bold",
        bold=True,
    )
    italic_run = InlineRunNode(
        node_type=NodeType.INLINE_RUN,
        id="test-2",
        text="italic",
        italic=True,
    )
    underline_run = InlineRunNode(
        node_type=NodeType.INLINE_RUN,
        id="test-3",
        text="underline",
        underline=True,
    )
    paragraph = Paragraph(
        id="test-para",
        runs=[
            TextRun(node_type=NodeType.TEXT_RUN, id="test-4", text="This is "),
            bold_run,
            TextRun(node_type=NodeType.TEXT_RUN, id="test-5", text=" and "),
            italic_run,
            TextRun(node_type=NodeType.TEXT_RUN, id="test-6", text=" and "),
            underline_run,
            TextRun(node_type=NodeType.TEXT_RUN, id="test-7", text=" text."),
        ],
    )
    doc = Document(id="test-doc", blocks=[paragraph])

    # Write to DOCX
    with tempfile.TemporaryDirectory() as tmpdir:
        output_path = Path(tmpdir) / "output.docx"
        writer = DocxWriter()
        writer.write(doc, output_path)

        # Verify DOCX content
        docx_doc = DocxDocument(str(output_path))
        para = docx_doc.paragraphs[0]
        runs = para.runs

        # Check bold
        bold_runs = [r for r in runs if r.bold]
        assert len(bold_runs) == 1, f"Expected 1 bold run, found {len(bold_runs)}"
        assert bold_runs[0].text == "bold"

        # Check italic
        italic_runs = [r for r in runs if r.italic]
        assert len(italic_runs) == 1, f"Expected 1 italic run, found {len(italic_runs)}"
        assert italic_runs[0].text == "italic"

        # Check underline
        underline_runs = [r for r in runs if r.underline]
        assert len(underline_runs) == 1, f"Expected 1 underline run, found {len(underline_runs)}"
        assert underline_runs[0].text == "underline"


def test_inline_paragraph_rendering() -> None:
    """Test full paragraph with multiple formatting styles."""
    # Create IR document with complex paragraph
    bold_italic_run = InlineRunNode(
        node_type=NodeType.INLINE_RUN,
        id="test-1",
        text="bold italic",
        bold=True,
        italic=True,
    )
    plain_text = TextRun(node_type=NodeType.TEXT_RUN, id="test-2", text=" plain text ")
    code_node = InlineCodeNode(
        node_type=NodeType.INLINE_CODE,
        id="test-3",
        code="code",
    )
    more_text = TextRun(node_type=NodeType.TEXT_RUN, id="test-4", text=" more text.")

    paragraph = Paragraph(
        id="test-para",
        runs=[bold_italic_run, plain_text, code_node, more_text],
    )
    doc = Document(id="test-doc", blocks=[paragraph])

    # Write to DOCX
    with tempfile.TemporaryDirectory() as tmpdir:
        output_path = Path(tmpdir) / "output.docx"
        writer = DocxWriter()
        writer.write(doc, output_path)

        # Verify DOCX content
        docx_doc = DocxDocument(str(output_path))
        para = docx_doc.paragraphs[0]
        runs = para.runs

        # Check combined bold+italic
        bold_italic_runs = [r for r in runs if r.bold and r.italic]
        assert len(bold_italic_runs) == 1, (
            f"Expected 1 bold+italic run, found {len(bold_italic_runs)}"
        )
        assert bold_italic_runs[0].text == "bold italic"

        # Check code font
        code_runs = [r for r in runs if r.text == "code"]
        assert len(code_runs) == 1
        font_name = code_runs[0].font.name
        assert font_name and "Courier" in font_name, f"Expected Courier font, got '{font_name}'"
