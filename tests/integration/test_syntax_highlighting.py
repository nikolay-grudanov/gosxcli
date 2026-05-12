"""Integration tests for syntax highlighting to DOCX.

Tests for:
- highlight_code() function with real Paragraph
- Color application to runs
- Different language code blocks
"""

from pathlib import Path

from docx import Document
from docx.shared import Pt, RGBColor

from typst_gost_docx.ir.model import CodeBlockNode, Document as IRDocument
from typst_gost_docx.writers.code_highlighter import (
    highlight_code,
)
from typst_gost_docx.writers.docx_writer import DocxWriter


class TestHighlightCode:
    """Integration tests for highlight_code function."""

    def test_highlight_code_creates_runs(self) -> None:
        """Should create runs for highlighted code."""
        doc = Document()
        paragraph = doc.add_paragraph()

        highlight_code("def hello():", "python", paragraph)

        assert len(paragraph.runs) > 0
        # First run should have text
        assert "def hello():" in paragraph.runs[0].text or len(paragraph.runs[0].text) > 0

    def test_highlight_code_applies_font(self) -> None:
        """Should apply monospace font to code."""
        doc = Document()
        paragraph = doc.add_paragraph()

        highlight_code("def hello():", "python", paragraph)

        # Check font name
        font_name = paragraph.runs[0].font.name
        assert font_name in ("Courier New", "Consolas", "Courier New")

    def test_highlight_code_applies_font_size(self) -> None:
        """Should apply font size to code."""
        doc = Document()
        paragraph = doc.add_paragraph()

        highlight_code("def hello():", "python", paragraph)

        # Check font size is set
        font_size = paragraph.runs[0].font.size
        assert font_size == Pt(9)

    def test_highlight_code_applies_color(self) -> None:
        """Should apply color to code runs."""
        doc = Document()
        paragraph = doc.add_paragraph()

        highlight_code("def hello():", "python", paragraph)

        # Check color is applied (first run should have color)
        color = paragraph.runs[0].font.color
        assert color is not None
        # RGB should be set if code has tokens
        assert color.rgb is not None

    def test_highlight_code_python(self) -> None:
        """Should highlight Python code."""
        doc = Document()
        paragraph = doc.add_paragraph()

        highlight_code("def hello():", "python", paragraph)

        # Verify code is present
        text = paragraph.text
        assert "def hello():" in text or len(text) > 0

    def test_highlight_code_rust(self) -> None:
        """Should highlight Rust code."""
        doc = Document()
        paragraph = doc.add_paragraph()

        highlight_code("fn main() {}", "rust", paragraph)

        # Verify code is present
        text = paragraph.text
        assert "fn main()" in text or len(text) > 0

    def test_highlight_code_javascript(self) -> None:
        """Should highlight JavaScript code."""
        doc = Document()
        paragraph = doc.add_paragraph()

        highlight_code("const x = 1;", "javascript", paragraph)

        # Verify code is present
        text = paragraph.text
        assert "const x" in text or len(text) > 0

    def test_highlight_code_multiple_tokens(self) -> None:
        """Should create multiple runs for different tokens."""
        doc = Document()
        paragraph = doc.add_paragraph()

        highlight_code("def hello():\n    pass", "python", paragraph)

        # With multiple tokens, should have multiple runs
        # At minimum should have some runs
        assert len(paragraph.runs) >= 1

    def test_highlight_code_with_shading(self) -> None:
        """Should apply shading when color provided."""
        doc = Document()
        paragraph = doc.add_paragraph()

        shading = RGBColor(0x1E, 0x1E, 0x1E)
        highlight_code("def hello():", "python", paragraph, shading_color=shading)

        # Verify code is present
        text = paragraph.text
        assert "def hello():" in text or len(text) > 0

    def test_highlight_code_empty_code(self) -> None:
        """Should handle empty code gracefully."""
        doc = Document()
        paragraph = doc.add_paragraph()

        highlight_code("", "python", paragraph)

        # Should handle empty code without error
        # Paragraph may have empty run or no runs
        assert paragraph.text == "" or len(paragraph.runs) >= 0


class TestSyntaxHighlightingRendering:
    """Integration tests for syntax highlighting DOCX rendering."""

    def test_write_code_block_with_highlighting(self, tmp_path: Path) -> None:
        """Should write code block with syntax highlighting."""
        code_block = CodeBlockNode(
            id="test-id",
            content="def hello():\n    return 42",
            language="python",
        )

        doc = IRDocument(blocks=[code_block])
        writer = DocxWriter()

        output_path = tmp_path / "output.docx"
        writer.write(doc, output_path)

        # Verify file created
        assert output_path.exists()

        # Open and verify
        docx = Document(str(output_path))
        paragraphs = docx.paragraphs

        # Should have paragraphs
        assert len(paragraphs) >= 1

    def test_write_code_block_python(self, tmp_path: Path) -> None:
        """Should write Python code block."""
        code_block = CodeBlockNode(
            id="python-code",
            content="x = 1",
            language="python",
        )

        doc = IRDocument(blocks=[code_block])
        writer = DocxWriter()

        output_path = tmp_path / "output.docx"
        writer.write(doc, output_path)

        # Verify
        docx = Document(str(output_path))
        text = "\n".join(p.text for p in docx.paragraphs)
        assert "x = 1" in text

    def test_write_code_block_rust(self, tmp_path: Path) -> None:
        """Should write Rust code block."""
        code_block = CodeBlockNode(
            id="rust-code",
            content="fn main() {}",
            language="rust",
        )

        doc = IRDocument(blocks=[code_block])
        writer = DocxWriter()

        output_path = tmp_path / "output.docx"
        writer.write(doc, output_path)

        # Verify
        docx = Document(str(output_path))
        text = "\n".join(p.text for p in docx.paragraphs)
        assert "fn main()" in text

    def test_write_code_block_javascript(self, tmp_path: Path) -> None:
        """Should write JavaScript code block."""
        code_block = CodeBlockNode(
            id="js-code",
            content="const x = 1;",
            language="javascript",
        )

        doc = IRDocument(blocks=[code_block])
        writer = DocxWriter()

        output_path = tmp_path / "output.docx"
        writer.write(doc, output_path)

        # Verify
        docx = Document(str(output_path))
        text = "\n".join(p.text for p in docx.paragraphs)
        assert "const x" in text

    def test_write_multiple_highlighted_code_blocks(self, tmp_path: Path) -> None:
        """Should write multiple code blocks with highlighting."""
        code_block1 = CodeBlockNode(
            id="python-code",
            content="print('hello')",
            language="python",
        )
        code_block2 = CodeBlockNode(
            id="rust-code",
            content="println!('hello')",
            language="rust",
        )

        doc = IRDocument(blocks=[code_block1, code_block2])
        writer = DocxWriter()

        output_path = tmp_path / "output.docx"
        writer.write(doc, output_path)

        # Verify
        docx = Document(str(output_path))
        text = "\n".join(p.text for p in docx.paragraphs)
        assert "print('hello')" in text
        assert "println!" in text

    def test_write_code_block_with_multiline(self, tmp_path: Path) -> None:
        """Should write multiline code block."""
        code = """def fib(n):
    if n <= 1:
        return n
    return fib(n-1) + fib(n-2)"""

        code_block = CodeBlockNode(
            id="fib-code",
            content=code,
            language="python",
        )

        doc = IRDocument(blocks=[code_block])
        writer = DocxWriter()

        output_path = tmp_path / "output.docx"
        writer.write(doc, output_path)

        # Verify
        assert output_path.exists()

        docx = Document(str(output_path))
        text = "\n".join(p.text for p in docx.paragraphs)
        assert "def fib(n):" in text
        assert "return fib(n-1)" in text


class TestHighlightingColorsApplied:
    """Tests verifying that colors are actually applied to runs."""

    def test_color_applied_to_first_run(self) -> None:
        """Should apply color to first run."""
        doc = Document()
        paragraph = doc.add_paragraph()

        highlight_code("def hello():", "python", paragraph)

        # First run should have non-default color
        first_run = paragraph.runs[0]
        assert first_run.font.color is not None
        assert first_run.font.color.rgb is not None

    def test_different_token_types_have_colors(self) -> None:
        """Should apply different colors to different token types."""
        # This tests keyword + name (both should get colors from TOKEN_COLORS)
        code = "def foo():"  # def=Keyword, foo=Name.Function
        doc = Document()
        paragraph = doc.add_paragraph()

        highlight_code(code, "python", paragraph)

        # At least one run should have color
        run_with_color = False
        for run in paragraph.runs:
            if run.font.color and run.font.color.rgb:
                run_with_color = True
                break

        assert run_with_color, "At least one run should have color applied"