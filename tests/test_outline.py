"""Tests for #outline() table of contents support."""

from typst_gost_docx.parser.scanner import TypstScanner
from typst_gost_docx.parser.extractor_v2 import TypstExtractorV2
from typst_gost_docx.ir.model import TOCNode, NodeType


def test_scanner_outline_token() -> None:
    """Test that scanner recognizes #outline( token."""
    text = "#outline()"
    scanner = TypstScanner(text)
    tokens = list(scanner.scan())

    outline_tokens = [t for t in tokens if t.type == "OUTLINE_START"]
    assert len(outline_tokens) == 1
    assert outline_tokens[0].value == "#outline("


def test_scanner_outline_with_title() -> None:
    """Test that scanner recognizes #outline( with title parameter."""
    text = '#outline(title: "Table of Contents")'
    scanner = TypstScanner(text)
    tokens = list(scanner.scan())

    outline_tokens = [t for t in tokens if t.type == "OUTLINE_START"]
    assert len(outline_tokens) == 1


def test_extractor_outline_basic() -> None:
    """Test that extractor extracts basic #outline()."""
    text = """
#outline()
= Introduction
Some content.
"""
    extractor = TypstExtractorV2(text, "test.typ")
    doc = extractor.extract()

    # Check that we have a TOC node
    toc_nodes = [node for node in doc.blocks if isinstance(node, TOCNode)]
    assert len(toc_nodes) == 1
    assert toc_nodes[0].title == "Содержание"
    assert toc_nodes[0].node_type == NodeType.TOC


def test_extractor_outline_with_custom_title() -> None:
    """Test that extractor extracts #outline() with custom title."""
    text = '''
#outline(title: "Table of Contents")
= Introduction
Some content.
'''
    extractor = TypstExtractorV2(text, "test.typ")
    doc = extractor.extract()

    # Check that we have a TOC node with custom title
    toc_nodes = [node for node in doc.blocks if isinstance(node, TOCNode)]
    assert len(toc_nodes) == 1
    assert toc_nodes[0].title == "Table of Contents"


def test_extractor_outline_with_headings() -> None:
    """Test that extractor extracts outline followed by headings."""
    text = """
#outline()
= Chapter 1
== Section 1.1
Content here.
= Chapter 2
"""
    extractor = TypstExtractorV2(text, "test.typ")
    doc = extractor.extract()

    # Check that we have TOC node and sections
    toc_nodes = [node for node in doc.blocks if isinstance(node, TOCNode)]
    assert len(toc_nodes) == 1

    from typst_gost_docx.ir.model import Section
    sections = [node for node in doc.blocks if isinstance(node, Section)]
    # We have 3 sections: Chapter 1 (level 1), Section 1.1 (level 2), Chapter 2 (level 1)
    assert len(sections) == 3
    assert sections[0].title[0].text == "Chapter 1"
    assert sections[0].level == 1
    assert sections[1].title[0].text == "Section 1.1"
    assert sections[1].level == 2
    assert sections[2].title[0].text == "Chapter 2"
    assert sections[2].level == 1


def test_docx_writer_outline(tmp_path) -> None:
    """Test that DocxWriter writes TOC node correctly."""
    from docx import Document
    from typst_gost_docx.writers.docx_writer import DocxWriter
    from typst_gost_docx.ir.model import Document as IRDocument, TOCNode

    # Create IR document with TOC
    ir_doc = IRDocument(id="test-doc")
    toc = TOCNode(id="test-toc", title="Table of Contents")
    ir_doc.blocks.append(toc)

    # Write to DOCX
    output_path = tmp_path / "test_toc.docx"
    writer = DocxWriter()
    _ = writer.write(ir_doc, output_path)

    # Verify file was created
    assert output_path.exists()

    # Verify TOC content
    doc = Document(str(output_path))
    paragraphs = doc.paragraphs

    # First paragraph should be the TOC title with Heading 1 style
    assert len(paragraphs) >= 2
    assert paragraphs[0].text == "Table of Contents"
    if paragraphs[0].style:
        assert paragraphs[0].style.name == "Heading 1"

    # Second paragraph should be the placeholder
    assert "Table of Contents will be generated here" in paragraphs[1].text


def test_multiple_outlines() -> None:
    """Test that document can have multiple outline() calls."""
    text = """
#outline()
= Chapter 1
Content.
#outline(title: "Table of Figures")
= Chapter 2
More content.
"""
    extractor = TypstExtractorV2(text, "test.typ")
    doc = extractor.extract()

    # Check that we have two TOC nodes
    toc_nodes = [node for node in doc.blocks if isinstance(node, TOCNode)]
    assert len(toc_nodes) == 2
    assert toc_nodes[0].title == "Содержание"
    assert toc_nodes[1].title == "Table of Figures"
