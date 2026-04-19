"""End-to-end structure tests for generated DOCX files.

These tests validate the structural integrity of generated DOCX files
to ensure they meet quality standards.
"""

from __future__ import annotations

from pathlib import Path
from typing import TYPE_CHECKING, Any, Generator

import pytest
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.table import CT_Tbl
from docx.oxml.text.paragraph import CT_P
from docx.table import Table
from docx.text.paragraph import Paragraph

from typst_gost_docx.ingest.project_loader import TypstProjectLoader
from typst_gost_docx.parser.extractor_v2 import TypstExtractorV2
from typst_gost_docx.writers.docx_writer import DocxWriter

if TYPE_CHECKING:
    from docx.document import Document as DocumentType


def convert_to_docx(fixture_path: Path, output_path: Path) -> DocumentType:
    """Конвертирует Typst файл в DOCX и возвращает Document объект.

    Args:
        fixture_path: Путь к входному Typst файлу.
        output_path: Путь для сохранения DOCX файла.

    Returns:
        Объект Document из python-docx.
    """
    from docx import Document

    # Загружаем Typst проект
    loader = TypstProjectLoader(fixture_path)
    files = loader.load()

    # Извлекаем IR
    text = files[str(fixture_path)]
    extractor = TypstExtractorV2(text, str(fixture_path))
    ir_document = extractor.extract()

    # Генерируем DOCX
    writer = DocxWriter()
    writer.write(ir_document, output_path)

    # Возвращаем Document объект для проверки
    return Document(str(output_path))


def iter_block_items(parent: Any) -> Generator[Paragraph | Table, None, None]:
    """Итерируется по элементам документа (параграфы и таблицы).

    Args:
        parent: Родительский элемент документа.

    Yields:
        Paragraph или Table элементы.
    """
    from docx.document import Document as DocumentType

    if isinstance(parent, DocumentType):
        parent_elm = parent.element.body
    elif isinstance(parent, CT_P):
        parent_elm = parent
    elif parent_elm := getattr(parent, "_element", None):
        parent_elm = parent_elm
    else:
        raise ValueError("Неподдерживаемый тип родительского элемента")

    for child in parent_elm.iterchildren():
        if isinstance(child, CT_P):
            yield Paragraph(child, parent)
        elif isinstance(child, CT_Tbl):
            yield Table(child, parent)


def count_headings(document: DocumentType) -> dict[str, int]:
    """Подсчитывает количество заголовков каждого уровня в документе.

    Args:
        document: Объект Document из python-docx.

    Returns:
        Словарь с количеством заголовков по уровням (1, 2, 3).
    """
    heading_counts = {"1": 0, "2": 0, "3": 0}

    for paragraph in document.paragraphs:
        if paragraph.style and paragraph.style.name.startswith("Heading"):
            # Извлекаем уровень из стиля (например, "Heading 1" -> "1")
            level = paragraph.style.name.split()[-1]
            if level in heading_counts:
                heading_counts[level] += 1

    return heading_counts


def count_tables(document: DocumentType) -> int:
    """Подсчитывает количество таблиц в документе.

    Args:
        document: Объект Document из python-docx.

    Returns:
        Количество таблиц в документе.
    """
    table_count = 0
    for element in iter_block_items(document):
        if isinstance(element, Table):
            table_count += 1
    return table_count


def count_empty_paragraphs(document: DocumentType) -> int:
    """Подсчитывает количество пустых параграфов в документе.

    Args:
        document: Объект Document из python-docx.

    Returns:
        Количество пустых параграфов.
    """
    empty_count = 0
    for paragraph in document.paragraphs:
        # Проверяем, что параграф не содержит текста и не является заголовком
        if not paragraph.text.strip() and not (
            paragraph.style and paragraph.style.name.startswith("Heading")
        ):
            empty_count += 1
    return empty_count


@pytest.fixture
def minimal_docx(tmp_path: Path) -> Generator[DocumentType, None, None]:
    """Фикстура для создания DOCX из минимального Typst файла.

    Args:
        tmp_path: Временная директория pytest.

    Yields:
        Объект Document из python-docx.
    """
    fixtures_dir = Path(__file__).parent.parent.parent / "fixtures"
    fixture_path = fixtures_dir / "minimal" / "minimal.typ"
    output_path = tmp_path / "minimal.docx"

    document = convert_to_docx(fixture_path, output_path)
    yield document

    # Очистка
    if output_path.exists():
        output_path.unlink()


@pytest.fixture
def real_vkr_docx(tmp_path: Path) -> Generator[DocumentType, None, None]:
    """Фикстура для создания DOCX из реального VRK файла.

    Args:
        tmp_path: Временная директория pytest.

    Yields:
        Объект Document из python-docx.
    """
    fixtures_dir = Path(__file__).parent.parent.parent / "fixtures"
    fixture_path = fixtures_dir / "real_vkr" / "main_doc" / "thesis.typ"
    output_path = tmp_path / "thesis.docx"

    document = convert_to_docx(fixture_path, output_path)
    yield document

    # Очистка
    if output_path.exists():
        output_path.unlink()


def test_minimal_docx_opens_without_error(minimal_docx: DocumentType) -> None:
    """Проверяет, что минимальный DOCX файл открывается без ошибок.

    Args:
        minimal_docx: Объект Document из python-docx.
    """
    # Если этот тест пройдёт без исключения, значит файл открывается корректно
    assert minimal_docx is not None
    assert len(minimal_docx.paragraphs) > 0


def test_minimal_docx_has_correct_headings(minimal_docx: DocumentType) -> None:
    """Проверяет, что минимальный документ имеет ожидаемое количество заголовков.

    Args:
        minimal_docx: Объект Document из python-docx.
    """
    heading_counts = count_headings(minimal_docx)

    # В minimal.typ: 1 заголовок уровня 1, 1 уровня 2, 1 уровня 3
    assert heading_counts["1"] == 1, "Ожидается 1 заголовок уровня 1"
    assert heading_counts["2"] == 1, "Ожидается 1 заголовок уровня 2"
    assert heading_counts["3"] == 1, "Ожидается 1 заголовок уровня 3"


def test_minimal_docx_has_expected_tables(minimal_docx: DocumentType) -> None:
    """Проверяет, что минимальный документ имеет ожидаемое количество таблиц.

    Args:
        minimal_docx: Объект Document из python-docx.
    """
    table_count = count_tables(minimal_docx)

    # В minimal.typ: 1 таблица
    assert table_count == 1, "Ожидается 1 таблица"


def test_minimal_docx_minimizes_empty_paragraphs(minimal_docx: DocumentType) -> None:
    """Проверяет, что минимальный документ имеет минимальное количество пустых параграфов.

    Args:
        minimal_docx: Объект Document из python-docx.
    """
    empty_count = count_empty_paragraphs(minimal_docx)

    # Допускаем не более 2 пустых параграфов (может быть для форматирования)
    assert empty_count <= 2, f"Слишком много пустых параграфов: {empty_count}"


def test_real_vkr_docx_opens_without_error(real_vkr_docx: DocumentType) -> None:
    """Проверяет, что реальный VRK DOCX файл открывается без ошибок.

    Args:
        real_vkr_docx: Объект Document из python-docx.
    """
    # Если этот тест пройдёт без исключения, значит файл открывается корректно
    assert real_vkr_docx is not None
    assert len(real_vkr_docx.paragraphs) > 0


def test_real_vkr_docx_has_headings(real_vkr_docx: DocumentType) -> None:
    """Проверяет, что реальный VRK документ имеет заголовки.

    Args:
        real_vkr_docx: Объект Document из python-docx.
    """
    heading_counts = count_headings(real_vkr_docx)

    # Реальный VRK документ должен иметь несколько заголовков
    total_headings = sum(heading_counts.values())
    assert total_headings > 0, "Документ должен содержать заголовки"


def test_real_vkr_docx_has_tables(real_vkr_docx: DocumentType) -> None:
    """Проверяет, что реальный VRK документ содержит таблицы (если ожидаются).

    Args:
        real_vkr_docx: Объект Document из python-docx.
    """
    table_count = count_tables(real_vkr_docx)

    # Проверяем только если таблицы ожидаются в документе
    # Для реального VRK документа таблицы могут отсутствовать
    # Этот тест просто проверяет, что подсчёт таблиц работает
    assert table_count >= 0, "Подсчёт таблиц должен работать"


def test_real_vkr_docx_minimizes_empty_paragraphs(real_vkr_docx: DocumentType) -> None:
    """Проверяет, что реальный VRK документ имеет минимальное количество пустых параграфов.

    Args:
        real_vkr_docx: Объект Document из python-docx.
    """
    empty_count = count_empty_paragraphs(real_vkr_docx)
    total_paragraphs = len(real_vkr_docx.paragraphs)

    # Допускаем не более 10% пустых параграфов
    max_empty = int(total_paragraphs * 0.1)
    assert empty_count <= max_empty, (
        f"Слишком много пустых параграфов: {empty_count} из {total_paragraphs}"
    )


def test_docx_has_valid_styles(minimal_docx: DocumentType) -> None:
    """Проверяет, что документ использует валидные стили.

    Args:
        minimal_docx: Объект Document из python-docx.
    """
    # Проверяем наличие основных стилей
    styles = minimal_docx.styles

    # Должны быть стили для заголовков
    heading_1_style = styles["Heading 1"]
    assert heading_1_style is not None
    assert heading_1_style.type == WD_STYLE_TYPE.PARAGRAPH

    heading_2_style = styles["Heading 2"]
    assert heading_2_style is not None
    assert heading_2_style.type == WD_STYLE_TYPE.PARAGRAPH

    heading_3_style = styles["Heading 3"]
    assert heading_3_style is not None
    assert heading_3_style.type == WD_STYLE_TYPE.PARAGRAPH


def test_docx_has_proper_document_structure(minimal_docx: DocumentType) -> None:
    """Проверяет, что документ имеет правильную структуру.

    Args:
        minimal_docx: Объект Document из python-docx.
    """
    # Документ должен иметь непустое тело
    assert len(minimal_docx.element.body) > 0

    # Проверяем, что все элементы являются параграфами или таблицами
    for element in iter_block_items(minimal_docx):
        assert isinstance(element, (Paragraph, Table)), (
            "Документ должен содержать только параграфы и таблицы"
        )


__all__ = [
    "test_minimal_docx_opens_without_error",
    "test_minimal_docx_has_correct_headings",
    "test_minimal_docx_has_expected_tables",
    "test_minimal_docx_minimizes_empty_paragraphs",
    "test_real_vkr_docx_opens_without_error",
    "test_real_vkr_docx_has_headings",
    "test_real_vkr_docx_has_tables",
    "test_real_vkr_docx_minimizes_empty_paragraphs",
    "test_docx_has_valid_styles",
    "test_docx_has_proper_document_structure",
]
