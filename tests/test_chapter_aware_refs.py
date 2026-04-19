"""Tests for chapter-aware numbering and cross-references."""

from typst_gost_docx.ir.model import (
    ChapterContext,
    CrossRefNode,
    Figure,
    TableNode,
    Equation,
    Section,
    Caption,
    NumberingKind,
)
from typst_gost_docx.parser.refs import RefResolver
from typst_gost_docx.writers.docx_writer import DocxWriter
from typst_gost_docx.config import RefLabels


def test_chapter_context_counters():
    """Test that ChapterContext tracks chapter and counters correctly."""
    ctx = ChapterContext()

    # Initial state
    assert ctx.chapter_number == 1
    assert ctx.figure_counter == 0
    assert ctx.table_counter == 0
    assert ctx.equation_counter == 0

    # Increment chapter
    ctx.chapter_number += 1
    assert ctx.chapter_number == 2

    # Increment counters
    ctx.figure_counter += 1
    assert ctx.figure_counter == 1

    ctx.table_counter += 1
    assert ctx.table_counter == 1

    # Reset counters on new chapter
    ctx.chapter_number += 1
    ctx.figure_counter = 0
    ctx.table_counter = 0
    ctx.equation_counter = 0

    assert ctx.chapter_number == 3
    assert ctx.figure_counter == 0
    assert ctx.table_counter == 0
    assert ctx.equation_counter == 0


def test_ref_resolver_infer_ref_kind():
    """Test that RefResolver infers ref kind from label prefixes."""
    resolver = RefResolver(cross_ref_map=None)

    # Test various prefixes
    assert resolver._infer_ref_kind("fig:results") == "fig"
    assert resolver._infer_ref_kind("tbl:data") == "tbl"
    assert resolver._infer_ref_kind("table:data") == "tbl"
    assert resolver._infer_ref_kind("eq:energy") == "eq"
    assert resolver._infer_ref_kind("equation:energy") == "eq"
    assert resolver._infer_ref_kind("ch:intro") == "ch"
    assert resolver._infer_ref_kind("chapter:intro") == "ch"
    assert resolver._infer_ref_kind("unknown:label") is None


def test_ref_resolver_cross_ref_node():
    """Test that RefResolver populates CrossRefNode correctly."""
    from typst_gost_docx.ir.model import CrossRefMap

    # Create a labeled figure
    figure = Figure(
        label="fig:results",
        number=2,
        chapter_number=1,
        caption=Caption(text="Results"),
    )

    # Create cross-ref map and register figure
    cross_ref_map = CrossRefMap()
    cross_ref_map.register("fig:results", figure)

    # Create resolver
    resolver = RefResolver(cross_ref_map)

    # Create CrossRefNode
    ref = CrossRefNode(target_label="fig:results")

    # Resolve
    resolver._resolve_cross_ref_node(ref)

    # Check that ref_kind, number, chapter_number are populated
    assert ref.ref_kind == "fig"
    assert ref.number == 2
    assert ref.chapter_number == 1


def test_docx_writer_format_cross_ref():
    """Test that DocxWriter formats cross-references correctly."""
    ref_labels = RefLabels(
        figure="Рис.",
        table="Таблица",
        equation="Формула",
    )

    writer = DocxWriter(ref_labels=ref_labels)

    # Test figure reference
    fig_ref = CrossRefNode(
        target_label="fig:results",
        ref_kind="fig",
        number=2,
        chapter_number=1,
    )
    fig_text = writer._format_cross_ref(fig_ref)
    assert fig_text == "Рис. 1.2"

    # Test table reference
    tbl_ref = CrossRefNode(
        target_label="tbl:data",
        ref_kind="tbl",
        number=1,
        chapter_number=2,
    )
    tbl_text = writer._format_cross_ref(tbl_ref)
    assert tbl_text == "Таблица 2.1"

    # Test equation reference
    eq_ref = CrossRefNode(
        target_label="eq:energy",
        ref_kind="eq",
        number=3,
        chapter_number=1,
    )
    eq_text = writer._format_cross_ref(eq_ref)
    assert eq_text == "Формула (1.3)"

    # Test reference without ref_kind (should infer)
    unk_ref = CrossRefNode(
        target_label="fig:unknown",
        number=5,
        chapter_number=3,
    )
    unk_text = writer._format_cross_ref(unk_ref)
    assert unk_text == "Рис. 3.5"


def test_docx_writer_write_caption():
    """Test that DocxWriter formats captions with chapter-aware numbering."""
    ref_labels = RefLabels(
        figure="Рис.",
        table="Табл.",
        equation="Формула",
    )

    writer = DocxWriter(ref_labels=ref_labels)

    # Mock document to test caption formatting without creating real DOCX
    from docx import Document

    writer.doc = Document()

    # Test figure caption
    fig_caption = Caption(
        text="Experimental results",
        numbering_kind=NumberingKind.FIGURE,
        number=2,
        chapter_number=1,
    )
    writer._write_caption(fig_caption)
    para = writer.doc.paragraphs[-1]
    assert para.text == "Рис. 1.2 — Experimental results"

    # Test table caption
    tbl_caption = Caption(
        text="Comparison of methods",
        numbering_kind=NumberingKind.TABLE,
        number=1,
        chapter_number=2,
    )
    writer._write_caption(tbl_caption)
    para = writer.doc.paragraphs[-1]
    assert para.text == "Табл. 2.1 — Comparison of methods"

    # Test equation caption
    eq_caption = Caption(
        text="Energy conservation",
        numbering_kind=NumberingKind.EQUATION,
        number=3,
        chapter_number=1,
    )
    writer._write_caption(eq_caption)
    para = writer.doc.paragraphs[-1]
    assert para.text == "Формула (1.3) — Energy conservation"


def test_docx_writer_infer_ref_kind():
    """Test that DocxWriter infers ref kind from label prefixes."""
    writer = DocxWriter()

    # Test various prefixes
    assert writer._infer_ref_kind("fig:results") == "fig"
    assert writer._infer_ref_kind("tbl:data") == "tbl"
    assert writer._infer_ref_kind("table:data") == "tbl"
    assert writer._infer_ref_kind("eq:energy") == "eq"
    assert writer._infer_ref_kind("equation:energy") == "eq"
    assert writer._infer_ref_kind("ch:intro") == "ch"
    assert writer._infer_ref_kind("chapter:intro") == "ch"
    assert writer._infer_ref_kind("unknown:label") is None


def test_figure_and_caption_numbering():
    """Test that figures get correct numbers and chapter numbers."""
    from docx import Document
    from typst_gost_docx.ir.model import TextRun

    ref_labels = RefLabels(figure="Рис.")
    writer = DocxWriter(ref_labels=ref_labels)
    writer.doc = Document()

    # Add first chapter
    sec1 = Section(level=1, title=[TextRun(text="Chapter 1")], label="ch:ch1")
    writer._write_section(sec1)

    # Add first figure
    fig1 = Figure(
        label="fig:results",
        caption=Caption(text="Results"),
    )
    writer._write_figure(fig1)

    assert fig1.number == 1
    assert fig1.chapter_number == 1
    assert fig1.caption.number == 1
    assert fig1.caption.chapter_number == 1

    # Add second figure
    fig2 = Figure(
        label="fig:comparison",
        caption=Caption(text="Comparison"),
    )
    writer._write_figure(fig2)

    assert fig2.number == 2
    assert fig2.chapter_number == 1
    assert fig2.caption.number == 2
    assert fig2.caption.chapter_number == 1

    # Add second chapter
    sec2 = Section(level=1, title=[TextRun(text="Chapter 2")], label="ch:ch2")
    writer._write_section(sec2)

    # Add figure in chapter 2
    fig3 = Figure(
        label="fig:chapter2",
        caption=Caption(text="Chapter 2 results"),
    )
    writer._write_figure(fig3)

    assert fig3.number == 1
    assert fig3.chapter_number == 2
    assert fig3.caption.number == 1
    assert fig3.caption.chapter_number == 2


def test_table_and_caption_numbering():
    """Test that tables get correct numbers and chapter numbers."""
    from docx import Document
    from typst_gost_docx.ir.model import ColSpec, TextRun

    ref_labels = RefLabels(table="Таблица")
    writer = DocxWriter(ref_labels=ref_labels)
    writer.doc = Document()

    # Add first chapter
    sec1 = Section(level=1, title=[TextRun(text="Chapter 1")], label="ch:ch1")
    writer._write_section(sec1)

    # Add first table
    tbl1 = TableNode(
        label="tbl:data",
        caption=Caption(text="Data"),
        columns=[ColSpec(), ColSpec()],
        rows=[],
    )
    writer._write_table(tbl1)

    assert tbl1.number == 1
    assert tbl1.chapter_number == 1
    assert tbl1.caption.number == 1
    assert tbl1.caption.chapter_number == 1

    # Add second table
    tbl2 = TableNode(
        label="tbl:comparison",
        caption=Caption(text="Comparison"),
        columns=[ColSpec(), ColSpec()],
        rows=[],
    )
    writer._write_table(tbl2)

    assert tbl2.number == 2
    assert tbl2.chapter_number == 1

    # Add second chapter
    sec2 = Section(level=1, title=[TextRun(text="Chapter 2")], label="ch:ch2")
    writer._write_section(sec2)

    # Add table in chapter 2
    tbl3 = TableNode(
        label="tbl:chapter2",
        caption=Caption(text="Chapter 2 data"),
        columns=[ColSpec(), ColSpec()],
        rows=[],
    )
    writer._write_table(tbl3)

    assert tbl3.number == 1
    assert tbl3.chapter_number == 2


def test_equation_and_caption_numbering():
    """Test that equations get correct numbers and chapter numbers."""
    from docx import Document
    from typst_gost_docx.ir.model import TextRun

    ref_labels = RefLabels(equation="Формула")
    writer = DocxWriter(ref_labels=ref_labels)
    writer.doc = Document()

    # Add first chapter
    sec1 = Section(level=1, title=[TextRun(text="Chapter 1")], label="ch:ch1")
    writer._write_section(sec1)

    # Add first equation
    eq1 = Equation(
        label="eq:energy",
        latex="E = mc^2",
        caption=Caption(text="Energy-mass equivalence"),
    )
    writer._write_equation(eq1)

    assert eq1.number == 1
    assert eq1.chapter_number == 1
    assert eq1.caption.number == 1
    assert eq1.caption.chapter_number == 1

    # Add second equation
    eq2 = Equation(
        label="eq:momentum",
        latex="p = mv",
        caption=Caption(text="Momentum"),
    )
    writer._write_equation(eq2)

    assert eq2.number == 2
    assert eq2.chapter_number == 1

    # Add second chapter
    sec2 = Section(level=1, title=[TextRun(text="Chapter 2")], label="ch:ch2")
    writer._write_section(sec2)

    # Add equation in chapter 2
    eq3 = Equation(
        label="eq:force",
        latex="F = ma",
        caption=Caption(text="Force"),
    )
    writer._write_equation(eq3)

    assert eq3.number == 1
    assert eq3.chapter_number == 2


def test_section_chapter_numbering():
    """Test that sections get correct chapter numbers."""
    from docx import Document
    from typst_gost_docx.ir.model import TextRun

    writer = DocxWriter()
    writer.doc = Document()

    # Initial state: chapter_number starts at 1
    assert writer.chapter_context.chapter_number == 1

    # Add first chapter
    sec1 = Section(level=1, title=[TextRun(text="Introduction")], label="ch:intro")
    writer._write_section(sec1)

    # First section should get chapter 1, context advances to chapter 2
    assert sec1.chapter_number == 1
    assert writer.chapter_context.chapter_number == 2  # Ready for next chapter
    assert writer.chapter_context.figure_counter == 0
    assert writer.chapter_context.table_counter == 0
    assert writer.chapter_context.equation_counter == 0

    # Add some figures
    fig1 = Figure(caption=Caption(text="Fig 1"))
    writer._write_figure(fig1)

    assert writer.chapter_context.figure_counter == 1

    # Add second chapter
    sec2 = Section(level=1, title=[TextRun(text="Methods")], label="ch:methods")
    writer._write_section(sec2)

    # Second section should get chapter 2, context advances to chapter 3
    assert sec2.chapter_number == 2
    assert writer.chapter_context.chapter_number == 3  # Ready for next chapter
    assert writer.chapter_context.figure_counter == 0  # Reset
    assert writer.chapter_context.table_counter == 0
    assert writer.chapter_context.equation_counter == 0

    # Add figure in chapter 2 (actually chapter 2, context is at 3)
    fig2 = Figure(caption=Caption(text="Fig 2"))
    writer._write_figure(fig2)

    # Figure should be in chapter 2 (the section's chapter), not 3
    assert fig2.number == 1
    assert fig2.chapter_number == 2
    assert writer.chapter_context.figure_counter == 1
