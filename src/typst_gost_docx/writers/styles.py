"""Resolution of GOST styles from DOCX template with non-standard style_id support."""

import logging
from importlib import resources
from pathlib import Path
from typing import Any, Optional, cast

from docx import Document
from docx.document import Document as _Document
from docx.styles.style import ParagraphStyle
from docx.styles.styles import Styles
from docx.text.paragraph import Paragraph as DocxParagraph

# NOTE: Iterative style lookup for templates with non-standard styleId.
# Replaces python-docx default __getitem__ with iteration-first approach
# that correctly handles both standard and non-standard styleId values
# (e.g., 781, 782, 783 for Heading 1-3 in ВКР template).
# Applied immediately at module load.

_APPLY_GUARD: bool = getattr(Styles, "__gosxcli_patched__", False)
_original_getitem = Styles.__getitem__


def _fixed_getitem(self: Styles, key: str) -> Any:
    """Fixed __getitem__ that iterates first to handle non-standard styleId.

    Args:
        self: Styles collection.
        key: Style name to lookup.

    Returns:
        Style object or result from original getitem.
    """
    # First: Try iteration - handles BOTH standard AND non-standard styleId
    for s in self:
        if s.name == key:
            return s

    # Second: Fall back to original python-docx logic for missing styles
    # This provides default styles (List Bullet, etc.) when using empty Document
    return _original_getitem(self, key)


if not _APPLY_GUARD:
    setattr(Styles, "__gosxcli_patched__", True)
    Styles.__getitem__ = _fixed_getitem  # type: ignore[method-assign]
    # keep guard variable for runtime reference


logger = logging.getLogger(__name__)


# IR element type → list of candidate style names (priority order)
IR_STYLE_MAP: dict[str, list[str]] = {
    "heading_1": ["Heading 1"],
    "heading_2": ["Heading 2"],
    "heading_3": ["Heading 3"],
    "heading_unnumbered": ["Заг_не_содержание"],
    "title": ["Титул"],
    "caption_figure": ["Подпись рисунков", "Caption"],
    "caption_table": ["Таблица название"],
    "equation": ["Формулы"],
    "normal": ["Normal"],
    "table_grid": ["Table Grid"],
    "list_bullet": ["List Bullet"],
    "list_number": ["List Number"],
}

# Keywords for detecting unnumbered headings (ВВЕДЕНИЕ, ЗАКЛЮЧЕНИЕ etc.)
UNNUMBERED_KEYWORDS: frozenset[str] = frozenset(
    {
        "РЕФЕРАТ",
        "ВВЕДЕНИЕ",
        "ЗАКЛЮЧЕНИЕ",
        "СПИСОК ЛИТЕРАТУРЫ",
        "ABSTRACT",
        "INTRODUCTION",
        "CONCLUSION",
        "BIBLIOGRAPHY",
        "REFERENCES",
    }
)


def get_default_template_path() -> Optional[Path]:
    """Получить путь к встроенному шаблону ГОСТ.

    Returns:
        Path к шаблону или None если шаблон недоступен.
    """
    # Template is in typst_gost_docx/templates/ package directory
    # Use importlib.resources to access it
    try:
        # Python 3.12+ approach
        ref = resources.files("typst_gost_docx.templates").joinpath(
            "Шаблон_оформления_ВКР_2026_новый.docx"
        )
        # resources.files returns a Traversable; convert to Path if possible
        with resources.as_file(ref) as path:
            if path.exists():
                return path
    except Exception:
        logger.warning("Встроенный шаблон ГОСТ недоступен")
    return None


def _clear_document_body(doc: _Document) -> None:
    """Очистить тело документа от контента, сохранив стили и секции.

    Удаляет все параграфы и таблицы из тела документа.
    Секции (page layout) сохраняются.

    Args:
        doc: Document для очистки.
    """
    from docx.oxml.ns import qn

    body = doc.element.body
    # Collect elements to remove (paragraphs and tables)
    # Keep sectPr (section properties) elements
    to_remove = []
    for child in body:
        tag = child.tag
        if tag == qn("w:p") or tag == qn("w:tbl"):
            to_remove.append(child)
    for element in to_remove:
        body.remove(element)


def load_document(reference_doc: Optional[Path] = None) -> _Document:
    """Загрузить Document из шаблона.

    Priority: reference_doc → built-in template → Document() with warnings.
    Template content (paragraphs, tables) is cleared; only styles are preserved.

    Args:
        reference_doc: Путь к пользовательскому шаблону (приоритет).

    Returns:
        Document загруженный из шаблона. Fallback на Document() если шаблон недоступен.
    """
    # Step 1: custom reference_doc provided and exists
    if reference_doc is not None:
        if reference_doc.exists():
            logger.info("Загрузка пользовательского шаблона: %s", reference_doc)
            doc = Document(str(reference_doc))
            _clear_document_body(doc)
            return doc
        else:
            logger.warning(
                "Пользовательский шаблон не найден: %s, используется встроенный",
                reference_doc,
            )

    # Step 2: try built-in template
    default_path = get_default_template_path()
    if default_path is not None:
        logger.info("Загрузка встроенного шаблона ГОСТ: %s", default_path)
        doc = Document(str(default_path))
        _clear_document_body(doc)
        return doc

    # Step 3: fallback to empty document with minimal GOST styles
    logger.warning("Шаблон недоступен, создаётся пустой документ с базовыми стилями")
    doc = Document()
    initialize_fallback_styles(doc)
    return doc


def is_unnumbered_heading(text: str) -> bool:
    """Определить, является ли заголовок ненумерованным по ключевому слову.

    Args:
        text: Текст заголовка.

    Returns:
        True если текст содержит ключевое слово ненумерованного заголовка.
    """
    return text.strip().upper() in UNNUMBERED_KEYWORDS


def initialize_fallback_styles(doc: _Document) -> None:
    """Инициализирует базовые стили для пустого документа.

    Создаёт минимальный набор стилей необходимых для GOST документа:
    - Normal: Times New Roman, 14pt
    - Heading 1, 2, 3: для структуры документа

    Args:
        doc: Document для инициализации стилей.
    """
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    from docx.shared import Pt, RGBColor

    # Set Normal style: Times New Roman, 14pt
    normal_style = doc.styles["Normal"]
    font = normal_style.font
    font.name = "Times New Roman"
    font.size = Pt(14)

    # Set East Asian and Complex Script fonts for Normal
    rPr = normal_style.element.get_or_add_rPr()
    rFonts = OxmlElement("w:rFonts")
    rFonts.set(qn("w:ascii"), "Times New Roman")
    rFonts.set(qn("w:hAnsi"), "Times New Roman")
    rFonts.set(qn("w:cs"), "Times New Roman")
    rPr.insert(0, rFonts)

    # Create Heading styles with black color, Times New Roman
    for level in range(1, 4):
        heading_style = doc.styles[f"Heading {level}"]
        heading_style.font.color.rgb = RGBColor(0, 0, 0)  # Black
        heading_style.font.name = "Times New Roman"

        # Set rFonts XML for heading styles
        h_rPr = heading_style.element.get_or_add_rPr()
        h_rFonts = OxmlElement("w:rFonts")
        h_rFonts.set(qn("w:ascii"), "Times New Roman")
        h_rFonts.set(qn("w:hAnsi"), "Times New Roman")
        h_rFonts.set(qn("w:cs"), "Times New Roman")
        h_rPr.insert(0, h_rFonts)


class StyleResolver:
    """Разрешение имён стилей из DOCX шаблона с учётом нестандартных style_id.

    Template has non-standard style_id values (781, 782, 783 for Heading 1-3),
    so direct doc.styles['Heading 1'] raises KeyError. This class handles
    both standard and non-standard style_id lookups with caching.
    """

    def __init__(self, doc: Optional[_Document]) -> None:
        self._doc = doc
        self._style_cache: dict[str, Optional[str]] = {}

    def resolve(self, ir_type: str) -> Optional[str]:
        """Разрешить IR-тип в имя стиля шаблона.

        Args:
            ir_type: Тип IR-элемента (например 'heading_1', 'caption_figure').

        Returns:
            Имя стиля из шаблона или None если стиль не найден.
            Всегда логирует warning при None.
        """
        candidates = IR_STYLE_MAP.get(ir_type)
        if candidates is None:
            logger.warning("Неизвестный IR-тип стиля: %s", ir_type)
            return None

        for candidate in candidates:
            # Check cache first
            if candidate in self._style_cache:
                cached = self._style_cache[candidate]
                if cached is not None:
                    return cached
                continue  # cached as not-found, try next candidate

            # Lookup in document
            found = self._lookup_style(candidate)
            if found is not None:
                self._style_cache[candidate] = found
                return found
            else:
                self._style_cache[candidate] = None  # mark as not-found

        # Fuzzy matching as last resort
        for candidate in candidates:
            fuzzy_result = self._resolve_by_fuzzy(candidate)
            if fuzzy_result is not None:
                self._style_cache[candidate] = fuzzy_result
                return fuzzy_result

        # Пробуем создать стиль динамически
        fallback_style = self._create_fallback_style(ir_type)
        if fallback_style is not None:
            return fallback_style

        logger.warning(
            "Стиль не найден для IR-типа '%s' (кандидаты: %s)",
            ir_type,
            candidates,
        )
        return None

    def _create_fallback_style(self, ir_type: str) -> Optional[str]:
        """Создать стиль динамически если он не найден в шаблоне.

        Создаёт базовые стили для списков и заголовков, которые могут
        отсутствовать в некоторых шаблонах (например, ВКР).

        Args:
            ir_type: IR тип элемента ('list_bullet', 'list_number', 'heading_N').

        Returns:
            Имя созданного стиля или None если не удалось создать.
        """
        from docx.enum.style import WD_STYLE_TYPE
        from docx.oxml import OxmlElement
        from docx.oxml.ns import qn
        from docx.shared import Pt, RGBColor

        if self._doc is None:
            return None

        try:
            if ir_type == "list_bullet":
                style = self._doc.styles.add_style("List Bullet", WD_STYLE_TYPE.PARAGRAPH)
                style.font.name = "Times New Roman"
                style.font.size = Pt(14)
                style.paragraph_format.left_indent = Pt(36)
                style.paragraph_format.first_line_indent = Pt(-18)
                logger.info("Создан fallback стиль: List Bullet")
                return "List Bullet"
            elif ir_type == "list_number":
                style = self._doc.styles.add_style("List Number", WD_STYLE_TYPE.PARAGRAPH)
                style.font.name = "Times New Roman"
                style.font.size = Pt(14)
                style.paragraph_format.left_indent = Pt(36)
                style.paragraph_format.first_line_indent = Pt(-18)
                logger.info("Создан fallback стиль: List Number")
                return "List Number"
            elif ir_type.startswith("heading_"):
                level = int(ir_type.split("_")[1])
                style_name = f"Heading {level}"
                style = self._doc.styles.add_style(style_name, WD_STYLE_TYPE.PARAGRAPH)
                style.font.name = "Times New Roman"
                style.font.bold = True
                sizes = {1: 16, 2: 15, 3: 14}
                style.font.size = Pt(sizes.get(level, 14))
                style.font.color.rgb = RGBColor(0, 0, 0)

                # Настраиваем rFonts
                rPr = style.element.get_or_add_rPr()
                rFonts = OxmlElement("w:rFonts")
                rFonts.set(qn("w:ascii"), "Times New Roman")
                rFonts.set(qn("w:hAnsi"), "Times New Roman")
                rFonts.set(qn("w:cs"), "Times New Roman")
                rPr.insert(0, rFonts)

                logger.info("Создан fallback стиль: %s", style_name)
                return style_name
            else:
                # Generic fallback to Normal
                return "Normal"
        except Exception as e:
            logger.warning("Не удалось создать fallback стиль %s: %s", ir_type, e)
            return None

    def apply_paragraph_style(self, paragraph: DocxParagraph, style_name: str) -> bool:
        """Применить стиль к параграфу с обходом бага python-docx.

        Работает даже когда python-docx не может разрешить имя стиля через [].
        Использует итеративный поиск и передаёт объект стиля вместо имени.

        Args:
            paragraph: Параграф python-docx.
            style_name: Имя стиля для применения.

        Returns:
            True если стиль применён, False если не найден.
        """
        if self._doc is None:
            logger.warning("Document не инициализирован, стиль не применён: %s", style_name)
            return False

        # Find style object via iteration (avoids BabelFish bug)
        style_obj = self._get_style_object(style_name)
        if style_obj is None:
            logger.warning("Стиль не найден для применения: %s", style_name)
            return False

        # Try style assignment (monkeypatch makes this work now)
        try:
            paragraph.style = style_name
            return True
        except KeyError:
            # Fallback to style object
            paragraph.style = cast(Any, style_obj)
            return True

    def _lookup_style(self, name: str) -> Optional[str]:
        """Найти стиль по имени: прямой lookup → итеративный fallback.

        Args:
            name: Имя стиля для поиска.

        Returns:
            Имя стиля или None если не найден.
        """
        if self._doc is None:
            return None

        # Direct lookup (works for standard style_id)
        try:
            style = self._doc.styles[name]
            return cast(str, style.name)
        except KeyError:
            pass

        # Iterative fallback (handles non-standard style_id like 781, 782, 783)
        for s in self._doc.styles:
            if s.name == name:
                return cast(str, s.name)

        return None

    def _get_style_object(self, name: str) -> Optional[ParagraphStyle]:
        """Получить объект стиля по имени (обход бага BabelFish).

        Использует итеративный поиск вместо словарного доступа,
        чтобы обойти баг python-docx с неправильным преобразованием имён.

        Args:
            name: Имя стиля для поиска.

        Returns:
            Объект стиля или None если не найден.
        """
        if self._doc is None:
            return None

        for s in self._doc.styles:
            typed_s = cast(ParagraphStyle, s)
            if typed_s.name == name:
                return typed_s

        return None

    def _resolve_by_fuzzy(self, name: str) -> Optional[str]:
        """Fuzzy matching стиля по подстроке.

        Args:
            name: Имя стиля для поиска.

        Returns:
            Найденное имя стиля или None.
        """
        if self._doc is None:
            return None

        name_lower = name.lower()
        for s in self._doc.styles:
            s_name = cast(str, s.name)
            if name_lower in s_name.lower() or s_name.lower() in name_lower:
                logger.info(
                    "Fuzzy match: '%s' → '%s'",
                    name,
                    s_name,
                )
                return s_name

        return None
