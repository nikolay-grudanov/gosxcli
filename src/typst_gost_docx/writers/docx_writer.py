"""DOCX writer for converting IR to Word documents."""

from pathlib import Path
from typing import Any, Optional, Sequence

from docx.document import Document as _Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.opc.constants import RELATIONSHIP_TYPE as RT
from docx.shared import Inches
from docx.text.paragraph import Paragraph as DocxParagraph
from ..ir.model import (
    Document as IRDocument,
    BaseNode,
    Section,
    Paragraph,
    ListBlock,
    TableNode,
    Figure,
    Caption,
    Equation,
    CrossReference,
    CrossRefNode,
    TOCNode,
    CitationNode,
    BibliographySection,
    BibliographyEntry,
    TextRun,
    InlineRunNode,
    InlineCodeNode,
    InlineMathNode,
    InlineLinkNode,
    CodeBlockNode,
    NumberingKind,
    CitationStyle,
    ChapterContext,
    ValidationResult,
)
from ..config import MathMode, RefLabels
from ..ir.validator import ReferenceValidator
from ..utils.ref_utils import infer_ref_kind
from .bookmarks import BookmarksManager
from .images import ImagesManager
from .tables import TablesManager
from .styles import StyleResolver, load_document, is_unnumbered_heading
from .code_highlighter import highlight_code, is_supported_language


class DocxWriter:
    def __init__(
        self,
        reference_doc: Optional[Path] = None,
        math_mode: MathMode = MathMode.FALLBACK,
        ref_labels: Optional[RefLabels] = None,
        bibliography_style: CitationStyle = CitationStyle.NUMERIC,
        base_dir: Optional[Path] = None,
    ):
        """Инициализирует DOCX writer.

        Args:
            reference_doc: Путь к шаблону DOCX для стилизации.
            math_mode: Режим рендеринга математических выражений.
            ref_labels: Локализованные метки для ссылок.
            bibliography_style: Стиль цитирования (numeric или author-year).
            base_dir: Базовая директория для резолвинга относительных путей (изображения и т.д.).
        """
        self.reference_doc = reference_doc
        self.math_mode = math_mode
        self.citation_style = bibliography_style
        self.entry_lookup: dict[str, BibliographyEntry] = {}
        self.doc: Optional[_Document] = None
        self.bookmarks_manager = BookmarksManager()
        self.images_manager = ImagesManager(base_dir=base_dir)
        self.tables_manager = TablesManager()
        from .lists import ListsManager

        self.lists_manager: Optional[ListsManager] = None
        self.chapter_context = ChapterContext()
        self.ref_labels = ref_labels or RefLabels()
        self.label_number_map: dict[str, tuple[int, int]] = {}  # label → (chapter_number, number)
        self._style_resolver: Optional[StyleResolver] = None

        self.stats = {
            "headings": 0,
            "paragraphs": 0,
            "tables": 0,
            "figures": 0,
            "equations": 0,
            "refs_resolved": 0,
            "refs_unresolved": 0,
            "warnings": 0,
        }

    @property
    def style_resolver(self) -> StyleResolver:
        """Lazy initialize StyleResolver for backward compatibility with tests."""
        if self._style_resolver is None:
            if self.doc is None:
                self.doc = load_document(self.reference_doc)
            self._style_resolver = StyleResolver(self.doc)
        assert self._style_resolver is not None
        return self._style_resolver

    @style_resolver.setter
    def style_resolver(self, value: Optional[StyleResolver]) -> None:
        self._style_resolver = value

    def _is_unnumbered_heading(self, text: str) -> bool:
        """Проверяет, является ли заголовок ненумерованным (ВВЕДЕНИЕ, ЗАКЛЮЧЕНИЕ и т.д.).

        Args:
            text: Текст заголовка для проверки.

        Returns:
            True если текст содержит ключевое слово ненумерованного заголовка.
        """
        return is_unnumbered_heading(text)

    def write(self, ir_document: IRDocument, output_path: Path) -> dict[str, Any]:
        self.doc = load_document(self.reference_doc)
        self.style_resolver = StyleResolver(self.doc)
        from .lists import ListsManager  # lazy: avoid circular import

        self.lists_manager = ListsManager(self.doc)

        self._write_document(ir_document)

        self.doc.save(str(output_path))
        return self.stats

    def _infer_ref_kind(self, label: str) -> Optional[str]:
        """Инферирует тип ссылки по префиксу метки.

        Args:
            label: Метка ссылки (например, fig:results).

        Returns:
            Тип ссылки (fig, tbl, eq) или None.
        """
        return infer_ref_kind(label)

    def _write_document(self, ir_doc: IRDocument) -> None:
        assert self.doc is not None, "Document not initialized"
        for block in ir_doc.blocks:
            self._write_block(block)

    def _write_block(self, block: BaseNode) -> None:
        if isinstance(block, Section):
            self._write_section(block)
        elif isinstance(block, Paragraph):
            self._write_paragraph(block)
        elif isinstance(block, ListBlock):
            self._write_list(block)
        elif isinstance(block, Figure):
            self._write_figure(block)
        elif isinstance(block, TableNode):
            self._write_table(block)
        elif isinstance(block, Equation):
            self._write_equation(block)
        elif isinstance(block, TOCNode):
            self._write_toc(block)
        elif isinstance(block, BibliographySection):
            self._write_bibliography(block)
        elif isinstance(block, CodeBlockNode):
            self._write_code_block(block)
        elif isinstance(block, CrossReference):
            pass

    def _write_section(self, section: Section) -> None:
        assert self.doc is not None, "Document not initialized"
        self.stats["headings"] += 1

        # Get title text early to detect unnumbered headings
        title_text = self._nodes_to_text(section.title) if section.title else ""
        is_unnumbered = self._is_unnumbered_heading(title_text) if title_text else False

        # Chapter-aware numbering: increment chapter on level 1 sections
        if section.level == 1:
            # Store current chapter number in the section
            section.chapter_number = self.chapter_context.chapter_number
            # Increment chapter number for next chapter
            self.chapter_context.chapter_number += 1
            # Reset counters when starting a new chapter
            self.chapter_context.figure_counter = 0
            self.chapter_context.table_counter = 0
            self.chapter_context.equation_counter = 0
            # Reset heading counters for new chapter
            self.chapter_context.heading_counters = [0, 0, 0]
        else:
            # Non-level-1 sections use current chapter number
            section.chapter_number = self.chapter_context.chapter_number

        section.number = self.chapter_context.section_counter

        if is_unnumbered:
            # Unnumbered headings: no number prefix, use special style
            number_str = ""
            heading_ir_type = "heading_unnumbered"
        else:
            # Issue 3: Sequential hierarchical heading numbering
            level_idx = section.level - 1  # 0-based index
            if 0 <= level_idx < 3:
                # Increment heading counter for this level
                self.chapter_context.heading_counters[level_idx] += 1
                # Reset all lower-level counters
                for i in range(level_idx + 1, 3):
                    self.chapter_context.heading_counters[i] = 0

                # Build number string: "1" for level 1, "1.1" for level 2, "1.1.1" for level 3
                nums = self.chapter_context.heading_counters[: level_idx + 1]
                number_str = ".".join(str(n) for n in nums)
            else:
                number_str = ""

            heading_ir_type = f"heading_{section.level}"

        if section.title:
            # Prepend number to title (only for numbered headings)
            if number_str:
                display_text = f"{number_str} {title_text}"
            else:
                display_text = title_text
            para = self.doc.add_paragraph(display_text)
            heading_style = self.style_resolver.resolve(heading_ir_type)
            if heading_style:
                self.style_resolver.apply_paragraph_style(para, heading_style)
            self.bookmarks_manager.add_bookmark_if_needed(para, section.label)

    def _write_paragraph(self, paragraph: Paragraph) -> None:
        assert self.doc is not None, "Document not initialized"
        self.stats["paragraphs"] += 1

        para = self.doc.add_paragraph()
        normal_style = self.style_resolver.resolve("normal")
        if normal_style:
            self.style_resolver.apply_paragraph_style(para, normal_style)
        self._write_inline_nodes(para, paragraph.runs)
        self.bookmarks_manager.add_bookmark_if_needed(para, paragraph.label)

    def _write_inline_math(self, para: DocxParagraph, latex: str) -> None:
        """Рендерит inline математику как OMML.

        Args:
            para: Параграф для добавления математики.
            latex: LaTeX строка (без обрамляющих $).
        """
        try:
            from .mml2omml import convert_mathml_to_omml
            from latex2mathml import converter

            mathml_str = converter.convert(latex)
            omml = convert_mathml_to_omml(mathml_str)
            if omml is not None:
                run = para.add_run()
                run._element.append(omml)
                return
        except Exception:
            pass
        # Fallback: render as plain text with $...$
        run = para.add_run(f"${latex}$")
        run.italic = True

    def _write_link(self, link: InlineLinkNode, para: DocxParagraph) -> None:
        """Emit an external hyperlink as ``<w:hyperlink r:id="...">``.

        The link's URL is registered as a relationship on the document
        part, then a hyperlinked run is appended to the paragraph with
        the standard blue-underline styling so it is visually
        recognisable as a link in Word and LibreOffice.
        """
        assert self.doc is not None, "Document not initialized"

        # Fallback: empty URL → just emit the visible text as plain run.
        if not link.url:
            if link.text:
                para.add_run(link.text)
            return

        # Register the URL as a relationship. ``relate_to`` with
        # ``is_external=True`` writes a TargetMode="External" entry
        # to word/_rels/document.xml.rels and returns a stable rId.
        r_id = self.doc.part.relate_to(link.url, RT.HYPERLINK, is_external=True)

        hyperlink = OxmlElement("w:hyperlink")
        hyperlink.set(qn("r:id"), r_id)

        run = OxmlElement("w:r")
        r_pr = OxmlElement("w:rPr")

        # Blue + underline so the link is visually obvious.
        color = OxmlElement("w:color")
        color.set(qn("w:val"), "0000FF")
        r_pr.append(color)
        underline = OxmlElement("w:u")
        underline.set(qn("w:val"), "single")
        r_pr.append(underline)

        run.append(r_pr)
        text_el = OxmlElement("w:t")
        text_el.text = link.text or link.url
        text_el.set(qn("xml:space"), "preserve")
        run.append(text_el)

        hyperlink.append(run)
        para._element.append(hyperlink)

        self.stats["links"] = self.stats.get("links", 0) + 1

    def _write_text_with_inline_math(self, para: DocxParagraph, text: str) -> None:
        """Разбивает текст на обычные фрагменты и inline math ($...$).

        Args:
            para: Параграф для добавления текста.
            text: Текст, возможно содержащий $...$ фрагменты.
        """
        import re

        parts = re.split(r"(\$[^$]+\$)", text)
        for part in parts:
            if part.startswith("$") and part.endswith("$") and len(part) > 2:
                latex = part[1:-1]
                self._write_inline_math(para, latex)
            elif part:
                para.add_run(part)

    def _write_inline_nodes(self, para: DocxParagraph, nodes: Sequence[BaseNode]) -> None:
        for node in nodes:
            if isinstance(node, TextRun):
                # Issue 4: Check for inline math ($...$) within text
                if "$" in node.text:
                    self._write_text_with_inline_math(para, node.text)
                else:
                    para.add_run(node.text)
            elif isinstance(node, InlineRunNode):
                run = para.add_run(node.text)
                if node.bold:
                    run.bold = True
                if node.italic:
                    run.italic = True
                if node.underline:
                    run.underline = True
            elif isinstance(node, InlineCodeNode):
                run = para.add_run(node.code)
                try:
                    run.style = "Code"
                except KeyError:
                    # Fallback: use Courier font if Code style doesn't exist
                    run.font.name = "Courier New"
            elif isinstance(node, InlineMathNode):
                self._write_inline_math(para, node.latex)
            elif isinstance(node, InlineLinkNode):
                self._write_link(node, para)
            elif isinstance(node, CrossReference):
                self._write_cross_reference(node, para)
            elif isinstance(node, CitationNode):
                self._write_citation(node, para)
            elif hasattr(node, "content"):
                self._write_inline_nodes(para, node.content)

    def _write_list(self, list_block: ListBlock) -> None:
        assert self.doc is not None, "Document not initialized"
        assert self.lists_manager is not None, "ListsManager not initialized"
        self.lists_manager.write_list(list_block)
        self.stats["lists"] = self.stats.get("lists", 0) + 1

    def _write_figure(self, figure: Figure) -> None:
        assert self.doc is not None, "Document not initialized"
        # Check if figure contains a table instead of an image
        if figure.table:
            # Treat table figure as a table (uses table counter)
            self._write_table(figure.table)
            # Update figure's caption number with table's number
            if figure.caption:
                figure.caption.number = figure.table.number
                figure.caption.chapter_number = figure.table.chapter_number
            # Write caption if present
            if figure.caption:
                self._write_caption(figure.caption, figure.label)
            return

        # Regular image figure
        self.stats["figures"] += 1

        # Increment figure counter and store in figure
        self.chapter_context.figure_counter += 1
        figure.number = self.chapter_context.figure_counter
        # Use current chapter number (chapter_context.chapter_number - 1 because we increment at section start)
        current_chapter = max(1, self.chapter_context.chapter_number - 1)
        figure.chapter_number = current_chapter

        # Update caption with number info
        if figure.caption:
            figure.caption.number = figure.number
            figure.caption.chapter_number = figure.chapter_number

        # Issue 5: Store label → (chapter, number) for cross-reference resolution
        if figure.label:
            self.label_number_map[figure.label] = (figure.chapter_number, figure.number)

        if figure.image_path:
            try:
                self.images_manager.add_image(self.doc, figure.image_path, width=Inches(5.0))
            except Exception:
                self.stats["warnings"] += 1

        if figure.caption:
            self._write_caption(figure.caption, figure.label)

    def _write_table(self, table: TableNode) -> None:
        self.stats["tables"] += 1

        # Increment table counter and store in table
        self.chapter_context.table_counter += 1
        table.number = self.chapter_context.table_counter
        # Use current chapter number (chapter_context.chapter_number - 1 because we increment at section start)
        current_chapter = max(1, self.chapter_context.chapter_number - 1)
        table.chapter_number = current_chapter

        # Update caption with number info
        if table.caption:
            table.caption.number = table.number
            table.caption.chapter_number = table.chapter_number

        # Issue 5: Store label → (chapter, number) for cross-reference resolution
        if table.label:
            self.label_number_map[table.label] = (table.chapter_number, table.number)

        self.tables_manager.add_table(self.doc, table)

    def _write_equation(self, equation: Equation) -> None:
        """Записывает математическое уравнение в документ.

        Использует math_mode для выбора метода рендеринга:
        - NATIVE: латекс → OMML через latex2mathml
        - IMAGE: заглушка с текстом (для MVP)
        - FALLBACK: сначала NATIVE, при ошибке → IMAGE

        Args:
            equation: IR узел уравнения.
        """
        assert self.doc is not None, "Document not initialized"
        self.stats["equations"] += 1

        # Increment equation counter and store in equation
        self.chapter_context.equation_counter += 1
        equation.number = self.chapter_context.equation_counter
        # Use current chapter number (chapter_context.chapter_number - 1 because we increment at section start)
        current_chapter = max(1, self.chapter_context.chapter_number - 1)
        equation.chapter_number = current_chapter

        # Update caption with number info
        if equation.caption:
            equation.caption.number = equation.number
            equation.caption.chapter_number = equation.chapter_number

        # Issue 5: Store label → (chapter, number) for cross-reference resolution
        if equation.label:
            self.label_number_map[equation.label] = (equation.chapter_number, equation.number)

        def _render_as_image(latex: str, doc: _Document) -> None:
            """Рендерит уравнение как текстовую заглушку (MVP)."""
            para = doc.add_paragraph()
            equation_style = self.style_resolver.resolve("equation")
            if equation_style:
                self.style_resolver.apply_paragraph_style(para, equation_style)
            para.add_run(f"[FORMULA: {latex}]")
            self.stats["warnings"] += 1

        def _render_as_native(latex: str, doc: _Document) -> bool:
            """Рендерит уравнение как OMML.

            Returns:
                True если успешно, False при ошибке.
            """
            import logging
            from .mml2omml import convert_mathml_to_omml
            from latex2mathml import converter

            logger = logging.getLogger(__name__)

            try:
                # Convert LaTeX to MathML, then to OMML
                mathml_str = converter.convert(latex)
                omml = convert_mathml_to_omml(mathml_str)

                if omml is None:
                    logger.warning(f"Failed to convert equation to OMML. Latex: {latex}")
                    return False

                para = doc.add_paragraph()
                equation_style = self.style_resolver.resolve("equation")
                if equation_style:
                    self.style_resolver.apply_paragraph_style(para, equation_style)
                run = para.add_run()
                run._element.append(omml)
                return True
            except Exception as e:
                logger.warning(f"Failed to render equation as OMML: {e}. Latex: {latex}")
                return False

        # Логика выбора режима рендеринга
        if self.math_mode == MathMode.IMAGE:
            _render_as_image(equation.latex, self.doc)
        elif self.math_mode == MathMode.NATIVE:
            success = _render_as_native(equation.latex, self.doc)
            if not success:
                self.stats["warnings"] += 1
        elif self.math_mode == MathMode.FALLBACK:
            success = _render_as_native(equation.latex, self.doc)
            if not success:
                _render_as_image(equation.latex, self.doc)

        if equation.caption:
            self._write_caption(equation.caption, equation.label)

    def _write_toc(self, toc: TOCNode) -> None:
        """Записывает оглавление в документ.

        Использует простой подход: заголовок + placeholder текст.
        python-docx имеет ограниченную поддержку TOC fields,
        поэтому в MVP используем placeholder.

        Args:
            toc: IR узел оглавления.
        """
        assert self.doc is not None, "Document not initialized"
        # Add TOC title as Heading 1
        toc_heading = self.doc.add_paragraph(toc.title)
        heading_style = self.style_resolver.resolve("heading_1")
        if heading_style:
            self.style_resolver.apply_paragraph_style(toc_heading, heading_style)

        # Add placeholder paragraph for TOC content
        # In a full implementation, this would be a TOC field
        # For MVP, we add a placeholder that indicates where TOC should be
        placeholder = self.doc.add_paragraph()
        normal_style = self.style_resolver.resolve("normal")
        if normal_style:
            self.style_resolver.apply_paragraph_style(placeholder, normal_style)
        placeholder.add_run(
            "[Table of Contents will be generated here - Right-click and select Update Field]"
        )
        placeholder.runs[0].italic = True
        placeholder.runs[0].font.color.rgb = None  # Use default color

    def _write_bibliography(self, bib_section: BibliographySection) -> None:
        """Записывает библиографическую секцию в документ.

        Генерирует секцию "Список литературы" с форматированием по ГОСТ 7.32-2017.

        Стили цитирования:
        - AUTHOR_YEAR: Записи сортируются по алфавиту (author → year), формат (Author, Year)
        - NUMERIC: Записи в порядке цитирования, формат [1], [2], [3]

        Форматирование по ГОСТ 7.32-2017:
        - Hanging indent: 1.25 см (первая строка с отступом влево, остальные вправо)
        - Заголовок: "Список литературы" (стиль Heading 1)

        Args:
            bib_section: IR узел библиографической секции.
        """
        from docx.shared import Cm
        from ..ir.model import CitationStyle

        assert self.doc is not None, "Document not initialized"

        # Add bibliography heading
        # Uses Heading 1 style for consistent document structure
        bib_heading = self.doc.add_paragraph(bib_section.heading)
        heading_style = self.style_resolver.resolve("heading_1")
        if heading_style:
            self.style_resolver.apply_paragraph_style(bib_heading, heading_style)

        # Get entries to write (sorted if AUTHOR_YEAR style)
        entries = bib_section.entries
        if bib_section.style == CitationStyle.AUTHOR_YEAR:
            # Sort alphabetically by author last name, then by year
            # Sorting is handled by _get_bibliography_sort_key
            entries = sorted(entries, key=self._get_bibliography_sort_key)

        # Write each bibliography entry with hanging indent
        for i, entry in enumerate(entries, start=1):
            # Format entry text based on citation style
            entry_text = self._format_bibliography_entry(entry, i, bib_section.style)
            para = self.doc.add_paragraph()
            normal_style = self.style_resolver.resolve("normal")
            if normal_style:
                self.style_resolver.apply_paragraph_style(para, normal_style)

            # Apply hanging indent (1.25cm as per ГОСТ 7.32-2017)
            # First line indent is negative (hanging), rest is positive
            # This creates the classic bibliography formatting where
            # the first line is left-aligned and subsequent lines are indented
            para.paragraph_format.first_line_indent = Cm(-1.25)
            para.paragraph_format.left_indent = Cm(1.25)

            # Add formatted entry text to paragraph
            para.add_run(entry_text)

    def _get_bibliography_sort_key(self, entry: BibliographyEntry) -> tuple[str, str]:
        """Получает ключ сортировки для библиографической записи.

        Используется для сортировки записей в AUTHOR_YEAR стиле.
        Ключ сортировки: (фамилия_автора, год)

        Логика сортировки:
        - Извлекает фамилию автора (первое слово перед запятой или пробелом)
        - Преобразует в lowercase для case-insensitive сортировки
        - Записи без автора помещаются в конец (last_name = "zzz")
        - Записи без года сортируются первыми (year = "0000")

        Args:
            entry: Библиографическая запись для сортировки.

        Returns:
            Кортеж из (фамилия_автора, год) для сортировки.
        """
        # Extract last name for sorting
        author = entry.author or ""
        # Author format may be "Иванов А.А." or "Ivanov A.A., Petrov B.B."
        # Take first word (last name) and convert to lowercase
        # This handles both single-author and multi-author entries
        last_name = author.split()[0].lower() if author else "zzz"
        # Use "zzz" for entries without author to place them at the end
        year = entry.year or "0000"
        # Use "0000" for entries without year to place them first among same author
        return (last_name, year)

    def _format_bibliography_entry(
        self, entry: BibliographyEntry, number: int, style: CitationStyle = CitationStyle.NUMERIC
    ) -> str:
        """Форматирует библиографическую запись по ГОСТ 7.32-2017.

        Форматы по ГОСТ 7.32-2017:

        **Статья (ARTICLE):**
        Автор А.А. Название // Журнал. — Год. — С. XX-XX.

        **Книга (BOOK):**
        Автор Б.Б. Название. — Город: Издательство, Год. — 256 с.

        **Материалы конференции (INPROCEEDINGS):**
        Автор В.В. Название // Труды конференции. — Год. — С. XX-XX.

        **Разное (MISC, TECHREPORT, etc.):**
        Автор. Название. — Год.

        Особенности форматирования:
        - Точки в конце полей удаляются (rstrip("."))
        - Поля соединяются через ". " (точка + пробел)
        - Для NUMERIC стиля добавляется номер: "1. Автор. Название..."
        - Для AUTHOR_YEAR стиля используется отдельный метод

        Args:
            entry: BibliographyEntry для форматирования.
            number: Порядковый номер записи (для NUMERIC стиля).
            style: Стиль цитирования (numeric или author-year).

        Returns:
            Отформатированная строка записи.
        """
        from ..ir.model import BibliographyType

        parts = []

        if entry.entry_type == BibliographyType.ARTICLE:
            # Article format: Автор А.А. Название // Журнал. — Год. — С. XX-XX.
            # Use placeholder for missing fields
            if entry.author:
                parts.append(entry.author.rstrip("."))
            else:
                parts.append("[Без автора]")
            if entry.title:
                parts.append(entry.title.rstrip("."))
            else:
                parts.append("[Без названия]")
            if entry.journal:
                parts.append(entry.journal.rstrip("."))
            if entry.year:
                parts.append(entry.year)
            if entry.pages:
                parts.append(f"С. {entry.pages}")

        elif entry.entry_type == BibliographyType.BOOK:
            # Book format: Автор Б.Б. Название. — Город: Издательство, Год. — 256 с.
            if entry.author:
                parts.append(entry.author.rstrip("."))
            else:
                parts.append("[Без автора]")
            if entry.title:
                parts.append(entry.title.rstrip("."))
            else:
                parts.append("[Без названия]")
            # Publisher may already include address (e.g., "Москва: Наука")
            if entry.publisher:
                pub = entry.publisher.rstrip(".")
                if entry.address and entry.address not in pub:
                    parts.append(f"{entry.address}: {pub}")
                else:
                    parts.append(pub)
            if entry.year:
                parts.append(entry.year)
            if entry.pages:
                parts.append(f"{entry.pages} с.")

        elif entry.entry_type == BibliographyType.INPROCEEDINGS:
            # Inproceedings format: Автор В.В. Название // Труды конференции. — Год. — С. XX-XX.
            if entry.author:
                parts.append(entry.author.rstrip("."))
            else:
                parts.append("[Без автора]")
            if entry.title:
                parts.append(entry.title.rstrip("."))
            else:
                parts.append("[Без названия]")
            if entry.booktitle:
                parts.append(entry.booktitle.rstrip("."))
            if entry.year:
                parts.append(entry.year)
            if entry.pages:
                parts.append(f"С. {entry.pages}")

        else:
            # Generic format for other types (MISC, TECHREPORT, PHDTHESIS, etc.)
            if entry.author:
                parts.append(entry.author.rstrip("."))
            else:
                parts.append("[Без автора]")
            if entry.title:
                parts.append(entry.title.rstrip("."))
            else:
                parts.append("[Без названия]")
            if entry.year:
                parts.append(entry.year)

        # Join parts with '. ' (period space) as per GOST
        entry_text = ". ".join(parts)

        # Format based on citation style
        if style == CitationStyle.AUTHOR_YEAR:
            # AUTHOR_YEAR style: Use separate formatting method
            return self._format_bibliography_entry_author_year(entry)
        else:
            # NUMERIC style: "1. Author Title..."
            return f"{number}. {entry_text}"

    def _format_bibliography_entry_author_year(self, entry: BibliographyEntry) -> str:
        """Форматирует библиографическую запись в стиле author-year (ГОСТ Р 7.0.5-2008).

        Форматы по ГОСТ Р 7.0.5-2008 (автор-год):

        **Статья (ARTICLE):**
        Автор А.А. (Год) Название // Журнал. — С. XX-XX.

        **Книга (BOOK):**
        Автор Б.Б. (Год) Название. — Город: Издательство. — 256 с.

        **Материалы конференции (INPROCEEDINGS):**
        Автор В.В. (Год) Название // Труды конференции. — С. XX-XX.

        **Разное (MISC, TECHREPORT, etc.):**
        Автор (Год) Название.

        Особенности форматирования:
        - Год в скобках после автора: (2023)
        - Для статей и материалов конференции используется "//" перед названием журнала/конференции
        - Поля соединяются через пробел
        - Не добавляется номер записи (в отличие от NUMERIC стиля)
        - Используется эм-даш (—) между частями

        Args:
            entry: BibliographyEntry для форматирования.

        Returns:
            Отформатированная строка записи без нумерации.
        """
        from ..ir.model import BibliographyType

        parts = []
        year = entry.year or ""

        if entry.entry_type == BibliographyType.ARTICLE:
            # Article format: Автор А.А. (Год) Название // Журнал. — С. XX-XX.
            if entry.author:
                parts.append(entry.author.rstrip("."))
            if year:
                parts.append(f"({year})")
            if entry.title:
                parts.append(entry.title.rstrip("."))
            else:
                parts.append("[Без названия]")
            if entry.journal:
                parts.append(f"// {entry.journal.rstrip('.')}")
            if entry.pages:
                parts.append(f"— С. {entry.pages}")

        elif entry.entry_type == BibliographyType.BOOK:
            # Book format: Автор Б.Б. (Год) Название. — Город: Издательство. — 256 с.
            if entry.author:
                parts.append(entry.author.rstrip("."))
            if year:
                parts.append(f"({year})")
            if entry.title:
                parts.append(entry.title.rstrip("."))
            else:
                parts.append("[Без названия]")
            if entry.publisher:
                pub = entry.publisher.rstrip(".")
                if entry.address and entry.address not in pub:
                    parts.append(f"— {entry.address}: {pub}")
                else:
                    parts.append(f"— {pub}")
            if entry.pages:
                parts.append(f"— {entry.pages} с.")

        elif entry.entry_type == BibliographyType.INPROCEEDINGS:
            # Inproceedings format: Автор В.В. (Год) Название // Труды конференции. — С. XX-XX.
            if entry.author:
                parts.append(entry.author.rstrip("."))
            if year:
                parts.append(f"({year})")
            if entry.title:
                parts.append(entry.title.rstrip("."))
            else:
                parts.append("[Без названия]")
            if entry.booktitle:
                parts.append(f"// {entry.booktitle.rstrip('.')}")
            if entry.pages:
                parts.append(f"— С. {entry.pages}")

        else:
            # Generic format for other types (MISC, TECHREPORT, PHDTHESIS, etc.)
            if entry.author:
                parts.append(entry.author.rstrip("."))
            if year:
                parts.append(f"({year})")
            if entry.title:
                parts.append(entry.title.rstrip("."))
            else:
                parts.append("[Без названия]")

        # Join parts with single space (author-year style uses space separator)
        return " ".join(parts)

    def _write_cross_reference(self, ref: CrossReference, para: DocxParagraph) -> None:
        """Write a cross-reference as a hyperlink.

        ``ref.ref_text`` is expected to be populated by ``RefResolver``
        before the writer runs (see ``cli.py``). As a fallback for legacy
        IR or test paths where the resolver did not run, we still consult
        ``label_number_map`` so we don't silently lose resolution.
        """
        ref_text = ref.ref_text or ""
        if not ref_text and ref.target_label in self.label_number_map:
            chapter_num, num = self.label_number_map[ref.target_label]
            ref.number = num
            ref.chapter_number = chapter_num
            template = {
                "fig": "Рис. {chapter}.{number}",
                "tbl": "Табл. {chapter}.{number}",
                "eq": "Формула {chapter}.{number}",
            }.get(ref.ref_kind or "")
            if template:
                ref_text = template.format(chapter=chapter_num, number=num)

        target = self.bookmarks_manager.get_bookmark(ref.target_label)

        if target and ref_text:
            hyperlink = OxmlElement("w:hyperlink")
            hyperlink.set(qn("w:anchor"), ref.target_label)

            run = OxmlElement("w:r")
            rPr = OxmlElement("w:rPr")

            wcolor = OxmlElement("w:color")
            wcolor.set(qn("w:val"), "0000FF")
            rPr.append(wcolor)

            wu = OxmlElement("w:u")
            wu.set(qn("w:val"), "single")
            rPr.append(wu)

            run.append(rPr)

            wt = OxmlElement("w:t")
            wt.text = ref_text
            run.append(wt)

            hyperlink.append(run)
            para._element.append(hyperlink)

            self.stats["refs_resolved"] += 1
        else:
            # No bookmark yet or unresolved: write the visible text plainly
            # so the reader still sees something. Use ref_text if we managed
            # to compute it, otherwise fall back to the raw label so the bug
            # stays visible in QA rather than silently disappearing.
            visible = ref_text or ref.target_label
            # ``run`` from the ``if`` branch above is an ``OxmlElement``
            # (lxml), and reusing the same name here would trip mypy's
            # no-redef check. Discard the new Run immediately — the text
            # is appended to the paragraph as a side-effect of add_run.
            _unused: Any = para.add_run(visible)
            del _unused
            self.stats["refs_unresolved"] += 1
            self.stats["warnings"] += 1

    def _write_citation(self, citation: CitationNode, para: DocxParagraph) -> None:
        """Записывает inline citation marker.

        Supports both NUMERIC and AUTHOR_YEAR citation styles:
        - NUMERIC: [1], [2], [3]
        - AUTHOR_YEAR: (Иванов, 2023), (Петров, 2022)

        Note: For NUMERIC style, brackets [] are captured as TEXT tokens
        by the scanner, so we only output the number here.

        Args:
            citation: CitationNode с ключом и номером.
            para: Paragraph для добавления маркера.
        """
        if self.citation_style == CitationStyle.AUTHOR_YEAR:
            entry = self.entry_lookup.get(citation.key)
            if entry and entry.author and entry.year:
                # Extract last name (first word before comma or first word)
                author_last = entry.author.split(",")[0].split()[0] if entry.author else ""
                citation_text = f"({author_last}, {entry.year})"
            else:
                # Fallback to numeric if no author/year
                citation_text = f"[{citation.number}]"
        else:
            # NUMERIC style: brackets are added by scanner
            citation_text = f"{citation.number}"

        run = para.add_run(citation_text)
        run.italic = True

    def _format_cross_ref(self, ref: CrossRefNode) -> str:
        """Форматирует текст перекрёстной ссылки с chapter-aware нумерацией.

        Args:
            ref: CrossRefNode для форматирования.

        Returns:
            Отформатированная строка ссылки (например, "Рис. 1.2", "Таблица 2.1").
        """
        # Determine ref kind from target label prefix if not set
        if not ref.ref_kind:
            ref.ref_kind = infer_ref_kind(ref.target_label)

        # Get localized label from config
        label = ""
        if ref.ref_kind == "fig":
            label = self.ref_labels.figure
        elif ref.ref_kind == "tbl":
            label = self.ref_labels.table
        elif ref.ref_kind == "eq":
            label = self.ref_labels.equation
        elif ref.ref_kind == "ch":
            label = self.ref_labels.section

        # Format based on ref kind
        if label:
            if ref.ref_kind == "eq":
                # Equations use parentheses: "Формула (1.3)"
                formatted = f"{label} ({ref.chapter_number}.{ref.number})"
            else:
                # Others use standard format: "Рис. 1.2", "Таблица 2.1"
                formatted = f"{label} {ref.chapter_number}.{ref.number}"
        else:
            # Fallback to label.number format
            formatted = f"{ref.chapter_number}.{ref.number}"

        return formatted

    def _write_caption(self, caption: Caption, label: Optional[str] = None) -> Optional[object]:
        """Write caption paragraph and optionally register a bookmark.

        Returns the paragraph if ``label`` was provided, else ``None``.
        """
        assert self.doc is not None, "Document not initialized"
        # Get localized prefix from config (rename to avoid shadowing ``label`` arg)
        prefix = ""
        if caption.numbering_kind == NumberingKind.FIGURE:
            prefix = self.ref_labels.figure
        elif caption.numbering_kind == NumberingKind.TABLE:
            prefix = self.ref_labels.table
        elif caption.numbering_kind == NumberingKind.EQUATION:
            prefix = self.ref_labels.equation

        # Format: "Рис. 1.2 - Caption text" or "Таблица 2.1 - Caption text"
        # For equations, use parentheses: "Формула (1.3)"
        if prefix:
            if caption.numbering_kind == NumberingKind.EQUATION:
                formatted = f"{prefix} ({caption.chapter_number}.{caption.number})"
            else:
                formatted = f"{prefix} {caption.chapter_number}.{caption.number}"
        else:
            formatted = f"{caption.chapter_number}.{caption.number}"

        # Add separator if there's text
        if caption.text:
            formatted = f"{formatted} — {caption.text}"  # em-dash for GOST style

        para = self.doc.add_paragraph()
        # Use different styles for figure and table captions
        if caption.numbering_kind == NumberingKind.TABLE:
            caption_ir_type = "caption_table"
        else:
            caption_ir_type = "caption_figure"
        caption_style = self.style_resolver.resolve(caption_ir_type)
        if caption_style:
            self.style_resolver.apply_paragraph_style(para, caption_style)
        para.add_run(formatted)

        # Register a bookmark so cross-references can hyperlink to this caption
        if label:
            self.bookmarks_manager.add_bookmark_if_needed(para, label)
            return para
        return None

    def _write_code_block(self, code_block: CodeBlockNode) -> None:
        """Записывает блок кода в документ.

        Использует моноширинный шрифт (Courier New/Consolas) для отображения кода.
        Применяет синтаксическое подсвечивание через Pygments, если язык определён.
        Сохраняет отступы и переносы строк. Экранирует XML спецсимволы.

        Args:
            code_block: IR узел блока кода.
        """
        from docx.oxml.ns import qn
        from docx.oxml import OxmlElement
        from docx.shared import Pt, RGBColor

        assert self.doc is not None, "Document not initialized"

        # Determine if we should apply syntax highlighting
        language = code_block.language or "plain"
        apply_highlighting = is_supported_language(language)

        # Background color for code blocks
        bg_color = RGBColor(0x1E, 0x1E, 0x1E)  # Dark gray (VS Code dark)

        # Process code line by line to preserve formatting
        lines = code_block.content.split("\n")

        for line in lines:
            # Create paragraph for each code line
            para = self.doc.add_paragraph(style="Normal")

            # Apply dark background shading
            shading_elm = OxmlElement("w:shd")
            shading_elm.set(qn("w:fill"), "1E1E1E")
            # Same python-docx / lxml ElementBase vs _Element mismatch as
            # in code_highlighter — see comment there.
            pPr: Any = para._element.get_or_add_pPr()
            pPr.insert_element_before(shading_elm, "w:spacing")

            # Escape XML special characters
            escaped_line = self._escape_xml_text(line)

            # Apply syntax highlighting or simple monospace
            if apply_highlighting:
                highlight_code(escaped_line, language, para, bg_color)
            else:
                # Simple monospace without highlighting
                run = para.add_run(escaped_line)
                run.font.name = "Courier New"
                run.font.size = Pt(9)
                # Light gray for plain text
                run.font.color.rgb = RGBColor(0xD4, 0xD4, 0xD4)

            # Remove interline spacing
            para.paragraph_format.space_before = Pt(0)
            para.paragraph_format.space_after = Pt(0)
            para.paragraph_format.line_spacing = 1.0

    def _escape_xml_text(self, text: str) -> str:
        """Экранирует XML спецсимволы для безопасного использования в DOCX.

        Экранирует символы: & → &amp;, < → &lt;, > → &gt;
        Это необходимо для корректного XML well-formed output.

        Args:
            text: Текст для экранирования.

        Returns:
            Экранированный текст, безопасный для XML/DOCX.
        """
        # Порядок важен: сначала & (чтобы не экранировать уже экранированные)
        text = text.replace("&", "&amp;")
        text = text.replace("<", "&lt;")
        text = text.replace(">", "&gt;")
        return text

    def _nodes_to_text(self, nodes: Sequence[BaseNode]) -> str:
        text_parts = []
        for node in nodes:
            if hasattr(node, "text"):
                text_parts.append(node.text)
            elif hasattr(node, "content"):
                text_parts.append(self._nodes_to_text(node.content))
        return "".join(text_parts)

    def validate_references(self, ir_document: IRDocument) -> ValidationResult:
        """Выполняет bidirectional валидацию ссылок и меток.

        Сравнивает определённые метки со ссылками и находит:
        - Неопределённые ссылки (@label без определения)
        - Неиспользуемые метки (<label> без ссылок)

        Args:
            ir_document: IR документ для валидации.

        Returns:
            ValidationResult с результатами валидации.
        """
        validator = ReferenceValidator()
        validator.collect_from_document(ir_document)
        result = validator.validate()

        # Обновляем статистику
        self.stats["refs_unresolved"] = len(result.undefined_refs)
        if result.unreferenced_labels:
            self.stats["warnings"] += len(result.unreferenced_labels)

        return result
