"""Tables manager for handling table creation with colspan/rowspan support."""

import logging
from typing import TYPE_CHECKING, Any, Optional

from docx.oxml import OxmlElement
from docx.oxml.ns import qn

from ..ir.model import (
    IRNode,
    TableNode as IRTable,
    TableCellNode,
    ColSpec,
)

if TYPE_CHECKING:
    pass

logger = logging.getLogger(__name__)


class TablesManager:
    """Tables manager for creating DOCX tables with advanced features.

    Supports:
    - Column width specification via ColSpec
    - Cell colspan and rowspan
    - Cell alignment
    - Cell background fill
    - Table borders
    """

    def __init__(self) -> None:
        """Initialize tables manager."""
        self.table_counter = 0

    def add_table(self, doc: Any, table: IRTable) -> None:
        """Add a table to the document with advanced features.

        Args:
            doc: Word document object.
            table: IR table node with columns, header, rows.
        """
        self.table_counter += 1

        if not table.rows and not table.header:
            logger.debug(f"Table {self.table_counter}: No rows or header, skipping")
            return

        # Calculate number of columns (considering colspan)
        num_cols = self._calculate_num_columns(table)
        if num_cols == 0:
            logger.debug(f"Table {self.table_counter}: No columns, skipping")
            return

        # Calculate number of rows (considering rowspan)
        num_rows = self._calculate_num_rows(table)
        if num_rows == 0:
            logger.debug(f"Table {self.table_counter}: No rows, skipping")
            return

        # Create table
        word_table = doc.add_table(rows=num_rows, cols=num_cols)
        word_table.style = "Table Grid"

        # Set column widths from ColSpec
        self._set_table_grid(word_table, table.columns, num_cols)

        # Track which cells are hidden due to rowspan
        # cell_map[row][col] = True if cell is visible, False if hidden
        cell_map: list[list[bool]] = [[True] * num_cols for _ in range(num_rows)]

        # Write header row
        row_idx = 0
        if table.has_header and table.header and table.header.cells:
            self._write_header_row(word_table, table.header.cells, row_idx, cell_map)
            row_idx += 1

        # Write data rows
        for table_row in table.rows:
            self._write_data_row(word_table, table_row, row_idx, cell_map)
            row_idx += 1

        logger.debug(f"Table {self.table_counter}: {num_rows}x{num_cols} table created")

    def _calculate_num_columns(self, table: IRTable) -> int:
        """Calculate number of columns in the table.

        Args:
            table: IR table node.

        Returns:
            Number of columns (considering ColSpec or max cells in row).
        """
        # Use column count from columns spec if available
        if table.columns:
            return len(table.columns)

        # Fallback to max cells in first row
        max_cols = 0
        if table.header and table.header.cells:
            max_cols = max(max_cols, len(table.header.cells))
        if table.rows and table.rows[0]:
            max_cols = max(max_cols, len(table.rows[0]))
        return max_cols

    def _calculate_num_rows(self, table: IRTable) -> int:
        """Calculate number of rows in the table.

        Args:
            table: IR table node.

        Returns:
            Number of rows (including header).
        """
        num_rows = 0
        if table.has_header and table.header:
            num_rows += 1
        num_rows += len(table.rows)
        return num_rows

    def _set_table_grid(
        self,
        word_table: Any,
        columns: list[ColSpec],
        num_cols: int,
    ) -> None:
        """Set column widths via tblGrid element.

        Args:
            word_table: Word table object.
            columns: List of column specifications.
            num_cols: Number of columns in the table.
        """
        tbl = word_table._tbl

        # Find or create tblGrid
        tbl_grid = tbl.tblGrid
        if tbl_grid is None:
            tbl_grid = OxmlElement("w:tblGrid")
            tbl.insert(0, tbl_grid)

        # Clear existing grid columns
        for grid_col in tbl_grid.xpath("./w:gridCol"):
            tbl_grid.remove(grid_col)

        # Add grid columns based on ColSpec
        for i in range(num_cols):
            grid_col = OxmlElement("w:gridCol")

            if i < len(columns):
                col_spec = columns[i]

                # Convert percentage width to EMU (1% = 360,000 EMU, max 100%)
                if col_spec.width_percent is not None:
                    width_emu = int(col_spec.width_percent * 360000)
                    # Clamp to valid range (0-100%)
                    width_emu = max(0, min(36000000, width_emu))
                    grid_col.set(qn("w:w"), str(width_emu))
                elif col_spec.width is not None:
                    # width is in points, convert to EMU (1 pt = 360,000 EMU)
                    width_emu = int(col_spec.width * 360000)
                    grid_col.set(qn("w:w"), str(width_emu))

            tbl_grid.append(grid_col)

    def _write_header_row(
        self,
        word_table: Any,
        cells: list[TableCellNode],
        row_idx: int,
        cell_map: list[list[bool]],
    ) -> None:
        """Write header row to table.

        Args:
            word_table: Word table object.
            cells: List of header cells.
            row_idx: Row index in the table.
            cell_map: Cell visibility map.
        """
        col_idx = 0
        for cell in cells:
            # Skip if this cell position is already taken (hidden by rowspan)
            while col_idx < len(cell_map[row_idx]) and not cell_map[row_idx][col_idx]:
                col_idx += 1

            if col_idx >= len(cell_map[row_idx]):
                break

            word_cell = word_table.rows[row_idx].cells[col_idx]

            # Set cell content (may contain nested tables)
            self._write_cell_content(word_cell, cell.content)

            # Apply cell attributes
            self._set_cell_colspan(word_cell, cell.colspan)
            self._set_cell_rowspan(word_cell, cell.rowspan, row_idx, col_idx, cell_map)
            self._set_cell_alignment(word_cell, cell.align or "center")

            # If cell has fill color, apply it
            if cell.fill:
                self._set_cell_fill(word_cell, cell.fill)

            # Mark hidden cells due to colspan
            for offset in range(1, cell.colspan):
                if col_idx + offset < len(cell_map[row_idx]):
                    cell_map[row_idx][col_idx + offset] = False

            # Move to next visible column
            col_idx += cell.colspan

    def _write_data_row(
        self,
        word_table: Any,
        row_cells: list[TableCellNode],
        row_idx: int,
        cell_map: list[list[bool]],
    ) -> None:
        """Write data row to table.

        Args:
            word_table: Word table object.
            row_cells: List of cells in the row.
            row_idx: Row index in the table.
            cell_map: Cell visibility map.
        """
        col_idx = 0
        for cell in row_cells:
            # Skip if this cell position is already taken (hidden by rowspan)
            while col_idx < len(cell_map[row_idx]) and not cell_map[row_idx][col_idx]:
                col_idx += 1

            if col_idx >= len(cell_map[row_idx]):
                break

            word_cell = word_table.rows[row_idx].cells[col_idx]

            # Set cell content (may contain nested tables)
            self._write_cell_content(word_cell, cell.content)

            # Apply cell attributes
            self._set_cell_colspan(word_cell, cell.colspan)
            self._set_cell_rowspan(word_cell, cell.rowspan, row_idx, col_idx, cell_map)
            self._set_cell_alignment(word_cell, cell.align or "left")

            # If cell has fill color, apply it
            if cell.fill:
                self._set_cell_fill(word_cell, cell.fill)

            # Apply default borders
            self._set_cell_borders(word_cell)

            # Mark hidden cells due to colspan
            for offset in range(1, cell.colspan):
                if col_idx + offset < len(cell_map[row_idx]):
                    cell_map[row_idx][col_idx + offset] = False

            # Move to next visible column
            col_idx += cell.colspan

    def _set_cell_colspan(self, cell: Any, colspan: int) -> None:
        """Set colspan for a cell using gridSpan element.

        Args:
            cell: Word table cell object.
            colspan: Number of columns to span.
        """
        if colspan <= 1:
            return

        tc = cell._tc
        tc_pr = tc.get_or_add_tcPr()

        # Remove existing gridSpan if any
        for grid_span in tc_pr.xpath("./w:gridSpan"):
            tc_pr.remove(grid_span)

        # Add new gridSpan
        grid_span = OxmlElement("w:gridSpan")
        grid_span.set(qn("w:val"), str(colspan))
        tc_pr.append(grid_span)

    def _set_cell_rowspan(
        self,
        cell: Any,
        rowspan: int,
        row_idx: int,
        col_idx: int,
        cell_map: list[list[bool]],
    ) -> None:
        """Set rowspan for a cell using vMerge element.

        Args:
            cell: Word table cell object.
            rowspan: Number of rows to span.
            row_idx: Current row index.
            col_idx: Current column index.
            cell_map: Cell visibility map (modified in place).
        """
        if rowspan <= 1:
            return

        tc = cell._tc
        tc_pr = tc.get_or_add_tcPr()

        # Set vMerge to "restart" for the first row
        v_merge = OxmlElement("w:vMerge")
        v_merge.set(qn("w:val"), "restart")
        tc_pr.append(v_merge)

        # Mark cells below as hidden and set vMerge to "continue"
        num_rows = len(cell_map)
        for r in range(row_idx + 1, min(row_idx + rowspan, num_rows)):
            if col_idx < len(cell_map[r]):
                cell_map[r][col_idx] = False

                # Get the cell in the next row and set vMerge to "continue"
                # Note: We need to find the actual cell at this position
                # This is tricky because we're using cell_map to track visibility
                # For now, we'll set vMerge when we actually write those cells
                pass

    def _set_cell_alignment(self, cell: Any, alignment: Optional[str]) -> None:
        """Set horizontal alignment for a cell.

        Args:
            cell: Word table cell object.
            alignment: Alignment value (left, center, right, justify).
        """
        if not alignment:
            return

        # Map alignment values to Word values
        alignment_map = {
            "left": "left",
            "center": "center",
            "right": "right",
            "justify": "both",
        }

        word_alignment = alignment_map.get(alignment.lower(), "left")

        # Set alignment via paragraph properties
        for paragraph in cell.paragraphs:
            p_pr = paragraph._p.get_or_add_pPr()
            jc = p_pr.jc
            if jc is None:
                jc = OxmlElement("w:jc")
                p_pr.append(jc)
            jc.set(qn("w:val"), word_alignment)

    def _set_cell_fill(self, cell: Any, fill_color: str) -> None:
        """Set background fill color for a cell.

        Args:
            cell: Word table cell object.
            fill_color: Color value (hex color like "DDDDDD" or named color).
        """
        tc = cell._tc
        tc_pr = tc.get_or_add_tcPr()

        # Remove existing shd if any
        for shd in tc_pr.xpath("./w:shd"):
            tc_pr.remove(shd)

        # Add new shading
        shd = OxmlElement("w:shd")
        shd.set(qn("w:fill"), fill_color)
        tc_pr.append(shd)

    def _set_cell_borders(self, cell: Any, border_width: float = 0.5) -> None:
        """Set borders for a cell.

        Args:
            cell: Word table cell object.
            border_width: Border width in points.
        """
        tc = cell._tc
        tc_pr = tc.get_or_add_tcPr()

        # Remove existing borders if any
        tc_borders = tc_pr.find(qn("w:tcBorders"))
        if tc_borders is not None:
            tc_pr.remove(tc_borders)

        # Create new borders
        tc_borders = OxmlElement("w:tcBorders")

        # Border size: 1/8 point units (e.g., 4 = 0.5pt)
        border_size = int(border_width * 8)

        for border_name in ["top", "left", "bottom", "right"]:
            border = OxmlElement(f"w:{border_name}")
            border.set(qn("w:val"), "single")
            border.set(qn("w:sz"), str(border_size))
            border.set(qn("w:space"), "0")
            border.set(qn("w:color"), "auto")
            tc_borders.append(border)

        tc_pr.append(tc_borders)

    def _write_cell_content(self, word_cell: Any, nodes: list[IRNode]) -> None:
        """Write cell content to Word cell.

        Handles plain text and nested tables.

        Args:
            word_cell: Word table cell object.
            nodes: List of IR content nodes.
        """
        # Check if any node is a TableNode (nested table)
        from ..ir.model import TableNode as IRTable

        table_nodes: list[IRTable] = [node for node in nodes if isinstance(node, IRTable)]

        if table_nodes:
            # Handle nested tables
            for table_node in table_nodes:
                self._add_nested_table(word_cell, table_node)
        else:
            # Handle plain text content
            text_parts: list[str] = []
            for node in nodes:
                if isinstance(node, IRTable):
                    # This should not happen since we already filtered out tables
                    continue
                if hasattr(node, "text"):
                    text_parts.append(node.text)
                elif hasattr(node, "code"):
                    text_parts.append(node.code)
                elif hasattr(node, "content"):
                    # Recursively handle nested content
                    if isinstance(node.content, list):
                        text_parts.append(self._cell_content_to_text(node.content))
                    else:
                        text_parts.append(str(node.content))
            word_cell.text = "".join(text_parts)

    def _add_nested_table(self, word_cell: Any, table: IRTable) -> None:
        """Add a nested table inside a Word cell using lxml.

        Args:
            word_cell: Word table cell object.
            table: IR table node to nest.
        """
        # Get the cell's tc element (OpenXML)
        tc = word_cell._tc

        # Create nested table element
        nested_tbl = OxmlElement("w:tbl")

        # Add basic table properties
        tbl_pr = OxmlElement("w:tblPr")
        tbl_style = OxmlElement("w:tblStyle")
        tbl_style.set(qn("w:val"), "TableGrid")
        tbl_pr.append(tbl_style)
        nested_tbl.append(tbl_pr)

        # Add table grid
        tbl_grid = OxmlElement("w:tblGrid")
        nested_tbl.append(tbl_grid)

        # Calculate number of columns
        num_cols = self._calculate_num_columns(table)
        for i in range(num_cols):
            grid_col = OxmlElement("w:gridCol")
            # Set default width (can be adjusted based on ColSpec)
            grid_col.set(qn("w:w"), "2267")  # ~0.5 inch in EMU
            tbl_grid.append(grid_col)

        # Add rows to nested table
        # Header row
        row_idx = 0
        if table.has_header and table.header and table.header.cells:
            self._add_nested_table_row(nested_tbl, table.header.cells, row_idx, True)
            row_idx += 1

        # Data rows
        for table_row in table.rows:
            self._add_nested_table_row(nested_tbl, table_row, row_idx, False)
            row_idx += 1

        # Append the nested table to the cell
        # Note: We need to add it to the cell's p (paragraph) or create a new p
        # For simplicity, we'll append it to the cell's content
        tc.append(nested_tbl)

    def _add_nested_table_row(
        self,
        nested_tbl: Any,
        cells: list[TableCellNode],
        row_idx: int,
        is_header: bool,
    ) -> None:
        """Add a row to nested table.

        Args:
            nested_tbl: OpenXML table element.
            cells: List of cells to add.
            row_idx: Row index (for cell references).
            is_header: Whether this is a header row.
        """
        tr = OxmlElement("w:tr")
        nested_tbl.append(tr)

        for cell in cells:
            tc = OxmlElement("w:tc")
            tr.append(tc)

            # Add cell properties
            tc_pr = OxmlElement("w:tcPr")
            tc.append(tc_pr)

            # Set colspan
            if cell.colspan > 1:
                grid_span = OxmlElement("w:gridSpan")
                grid_span.set(qn("w:val"), str(cell.colspan))
                tc_pr.append(grid_span)

            # Set borders
            self._add_nested_cell_borders(tc_pr)

            # Add cell content (paragraph)
            p = OxmlElement("w:p")
            tc.append(p)

            # Add paragraph properties (alignment)
            p_pr = OxmlElement("w:pPr")
            p.append(p_pr)

            alignment = "center" if is_header else (cell.align or "left")
            jc = OxmlElement("w:jc")
            jc.set(qn("w:val"), alignment)
            p_pr.append(jc)

            # Add run with text
            r = OxmlElement("w:r")
            p.append(r)

            # Add text
            cell_text = self._cell_content_to_text(cell.content)
            t = OxmlElement("w:t")
            t.set(qn("xml:space"), "preserve")
            t.text = cell_text
            r.append(t)

    def _add_nested_cell_borders(self, tc_pr: Any) -> None:
        """Add borders to nested table cell.

        Args:
            tc_pr: OpenXML tcPr element.
        """
        tc_borders = OxmlElement("w:tcBorders")
        tc_pr.append(tc_borders)

        for border_name in ["top", "left", "bottom", "right"]:
            border = OxmlElement(f"w:{border_name}")
            border.set(qn("w:val"), "single")
            border.set(qn("w:sz"), "4")  # 0.5pt
            border.set(qn("w:space"), "0")
            border.set(qn("w:color"), "auto")
            tc_borders.append(border)

    def _cell_content_to_text(self, nodes: list[IRNode]) -> str:
        """Convert cell content nodes to plain text.

        Args:
            nodes: List of IR content nodes.

        Returns:
            Plain text representation.
        """
        text_parts = []
        for node in nodes:
            if hasattr(node, "text"):
                text_parts.append(node.text)
            elif hasattr(node, "code"):
                text_parts.append(node.code)
            elif hasattr(node, "content"):
                # Recursively handle nested content
                if isinstance(node.content, list):
                    text_parts.append(self._cell_content_to_text(node.content))
                else:
                    text_parts.append(str(node.content))
        return "".join(text_parts)
